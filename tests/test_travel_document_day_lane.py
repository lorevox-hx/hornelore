"""WO-TRAVEL-DOC-CLOSEOUT-01 — the timeline snapshot.

THE PRODUCT RULE THIS SUITE GUARDS
----------------------------------
    The visible trip timeline is the editable source of truth.
    Export Travel Document produces a DOCX snapshot of that timeline.

So the document is not a curated memoir and there is no approval gate
on it. Whatever an operator can see on the timeline is what the Word
file contains, in the same day and the same order.

WHY IT EXISTS
-------------
The Bismarck export of 2026-08-06 came out with Part I — "The Journey
in Order" — empty while the day cards held "Santa Fe to Bismarck",
"Downtown Bismarck" and "Radisson Hotel on Main Street". A transitive
read of `trip_memoir_preview`'s call graph found the cause: it reached
trips, trip_regions, trip_stops, trip_themes, trip_location_notes,
trip_sources, trip_photo_links and photos, and never trip_days. The
same export printed "Unplaced" over two photographs the operator had
placed on days by hand. 463 tests were green throughout; none of them
had a day row.

[An earlier version of this file tested an approval-gated design —
`trip_days.include_in_memoir` deciding which days reached the
document. Chris rejected that design before it was committed. The
migration stays applied so fresh installations match his database and
the column is DORMANT; a test here asserts that nothing on the export
path reads it.]

Run:
    PYTHONPYCACHEPREFIX=/tmp/pyc PYTHONPATH=server/code \\
        .venv/bin/python -m unittest tests.test_travel_document_day_lane
"""
from __future__ import annotations

import io
import os
import sqlite3
import sys
import tempfile
import unittest
import uuid
from pathlib import Path

_REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_REPO / "server" / "code"))

try:
    from docx import Document          # noqa: F401
    _READY, _WHY = True, ""
except Exception as exc:               # pragma: no cover - env dependent
    _READY, _WHY = False, str(exc)

TRIP = "trip-timeline"
PERSON = "person-timeline"


def _mk_db(path: str) -> None:
    """The real schema the timeline touches, hand-written so the columns
    it depends on are visible in the test rather than implied by 42
    migration files."""
    con = sqlite3.connect(path)
    con.executescript(
        """
        CREATE TABLE people (
            id TEXT PRIMARY KEY, display_name TEXT);
        CREATE TABLE turns (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            role TEXT, content TEXT, ts TEXT);
        CREATE TABLE trips (
            id TEXT PRIMARY KEY, person_id TEXT, title TEXT,
            start_date TEXT, end_date TEXT, summary TEXT,
            status TEXT, created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_regions (
            id TEXT PRIMARY KEY, trip_id TEXT, title TEXT,
            country_or_area TEXT, start_date TEXT, end_date TEXT,
            base_address TEXT, summary TEXT, ord INTEGER,
            created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_stops (
            id TEXT PRIMARY KEY, trip_id TEXT, trip_region_id TEXT,
            parent_trip_stop_id TEXT, location_name TEXT, title TEXT,
            stop_type TEXT, date_start TEXT, date_end TEXT, notes TEXT,
            thematic_tags_json TEXT, ord INTEGER,
            created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_themes (
            id TEXT PRIMARY KEY, trip_id TEXT, title TEXT, tag TEXT,
            description TEXT, ord INTEGER,
            created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_days (
            id TEXT PRIMARY KEY, trip_id TEXT, day_index INTEGER,
            date TEXT, title TEXT, main_location TEXT, lodging_base TEXT,
            trip_region_id TEXT, trip_stop_id TEXT,
            morning_notes TEXT, afternoon_notes TEXT, evening_notes TEXT,
            places_visited_json TEXT, meals_json TEXT,
            created_at TEXT, updated_at TEXT, reconcile_status TEXT,
            include_in_memoir INTEGER NOT NULL DEFAULT 0);
        CREATE TABLE trip_location_notes (
            id TEXT PRIMARY KEY, trip_id TEXT, trip_region_id TEXT,
            trip_stop_id TEXT, trip_day_id TEXT, note_title TEXT,
            note_text TEXT, source_type TEXT,
            include_in_memoir INTEGER DEFAULT 0, hidden INTEGER DEFAULT 0,
            source_surface TEXT, ord INTEGER DEFAULT 0,
            created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_sources (
            id TEXT PRIMARY KEY, trip_id TEXT, trip_region_id TEXT,
            trip_stop_id TEXT, trip_day_id TEXT, title TEXT,
            source_type TEXT, summary TEXT, pasted_text TEXT,
            link_url TEXT, filename TEXT, source_date TEXT,
            include_in_memoir INTEGER DEFAULT 0, hidden INTEGER DEFAULT 0,
            ord INTEGER DEFAULT 0, created_at TEXT, updated_at TEXT);
        CREATE TABLE photos (
            id TEXT PRIMARY KEY, narrator_id TEXT, image_path TEXT,
            description TEXT, date_value TEXT, narrator_ready INTEGER,
            caption_approved_for_lori INTEGER DEFAULT 0,
            operator_context_note TEXT, deleted_at TEXT);
        CREATE TABLE trip_photo_links (
            id TEXT PRIMARY KEY, trip_id TEXT, photo_id TEXT,
            trip_region_id TEXT, trip_stop_id TEXT, trip_day_id TEXT,
            narrator_caption TEXT, caption TEXT, taken_at TEXT,
            ord INTEGER, assignment_method TEXT,
            include_in_memoir INTEGER DEFAULT 0, hidden INTEGER DEFAULT 0,
            created_at TEXT, updated_at TEXT);
        CREATE TABLE trip_turn_links (
            id TEXT PRIMARY KEY, trip_id TEXT, trip_day_id TEXT,
            conv_id TEXT, user_turn_row_id INTEGER,
            assistant_turn_row_id INTEGER, captured_at TEXT,
            placement_source TEXT, placement_status TEXT,
            created_at TEXT, updated_at TEXT);
        """
    )
    con.execute("INSERT INTO people (id, display_name) VALUES (?,?)",
                (PERSON, "Christopher Todd Horne"))
    con.execute(
        "INSERT INTO trips (id, person_id, title, start_date, end_date,"
        " summary, status, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?)",
        (TRIP, PERSON, "Bismarck Trip", "2026-07-14", "2026-07-19",
         "A week in North Dakota.", "draft", "t", "t"))
    con.commit()
    con.close()


def _png(path: Path) -> Path:
    path.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753"
        "de0000000c4944415478da63f8ffff3f0005fe02fea735c9ab0000000049454e"
        "44ae426082"))
    return path


class TimelineTestBase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        if not _READY:                 # pragma: no cover
            raise unittest.SkipTest(f"ENV-SKIP: {_WHY}")

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.db = os.path.join(self.tmp.name, "t.sqlite3")
        _mk_db(self.db)
        import api.db as db
        from api.services import trip_repository as R
        self._old = db.DB_PATH
        db.DB_PATH = self.db
        if getattr(R, "db", None) is not None:
            R.db.DB_PATH = self.db
        self.R = R
        self.con = sqlite3.connect(self.db)

    def tearDown(self):
        self.con.close()
        import api.db as db
        from api.services import trip_repository as R
        db.DB_PATH = self._old
        if getattr(R, "db", None) is not None:
            R.db.DB_PATH = self._old
        self.tmp.cleanup()

    # ── fixture helpers ──────────────────────────────────────────────
    def add_day(self, index, date, **fields):
        did = f"day-{index}"
        cols = dict(id=did, trip_id=TRIP, day_index=index, date=date,
                    created_at="t", updated_at="t")
        cols.update(fields)
        keys = ", ".join(cols)
        marks = ", ".join("?" for _ in cols)
        self.con.execute(f"INSERT INTO trip_days ({keys}) VALUES ({marks})",
                         tuple(cols.values()))
        self.con.commit()
        return did

    def add_note(self, day_id=None, stop_id=None, text="a note",
                 title=None, hidden=0, approved=0):
        nid = str(uuid.uuid4())
        self.con.execute(
            "INSERT INTO trip_location_notes (id, trip_id, trip_day_id,"
            " trip_stop_id, note_title, note_text, include_in_memoir,"
            " hidden, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (nid, TRIP, day_id, stop_id, title, text, approved, hidden,
             "t", "t"))
        self.con.commit()
        return nid

    def add_source(self, day_id=None, title="a source", hidden=0,
                   approved=0):
        sid = str(uuid.uuid4())
        self.con.execute(
            "INSERT INTO trip_sources (id, trip_id, trip_day_id, title,"
            " source_type, include_in_memoir, hidden, created_at,"
            " updated_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (sid, TRIP, day_id, title, "itinerary", approved, hidden,
             "t", "t"))
        self.con.commit()
        return sid

    def add_photo(self, day_id=None, stop_id=None, approved=0,
                  narrator_caption=None, operator_caption=None,
                  machine_description=None, taken_at="2026-07-14",
                  hidden=0, deleted=None, on_disk=True):
        pid, lid = str(uuid.uuid4()), str(uuid.uuid4())
        path = str(_png(Path(self.tmp.name) / f"{pid}.png")) if on_disk \
            else os.path.join(self.tmp.name, "missing.png")
        self.con.execute(
            "INSERT INTO photos (id, narrator_id, image_path, description,"
            " narrator_ready, deleted_at) VALUES (?,?,?,?,1,?)",
            (pid, PERSON, path, machine_description, deleted))
        self.con.execute(
            "INSERT INTO trip_photo_links (id, trip_id, photo_id,"
            " trip_stop_id, trip_day_id, narrator_caption, caption,"
            " taken_at, ord, assignment_method, include_in_memoir,"
            " hidden, created_at, updated_at)"
            " VALUES (?,?,?,?,?,?,?,?,0,'operator',?,?,'t','t')",
            (lid, TRIP, pid, stop_id, day_id, narrator_caption,
             operator_caption, taken_at, approved, hidden))
        self.con.commit()
        return pid, lid

    def add_conversation(self, day_id=None, said="I walked to the river.",
                         replied="Tell me more about that.",
                         status="confirmed", at="2026-07-14T10:00:00Z"):
        u = self.con.execute(
            "INSERT INTO turns (role, content, ts) VALUES ('user',?,?)",
            (said, at)).lastrowid
        a = self.con.execute(
            "INSERT INTO turns (role, content, ts) VALUES ('assistant',?,?)",
            (replied, at)).lastrowid
        self.con.execute(
            "INSERT INTO trip_turn_links (id, trip_id, trip_day_id, conv_id,"
            " user_turn_row_id, assistant_turn_row_id, captured_at,"
            " placement_source, placement_status, created_at, updated_at)"
            " VALUES (?,?,?,'c1',?,?,?,'active_trip_day',?,'t','t')",
            (str(uuid.uuid4()), TRIP, day_id, u, a, at, status))
        self.con.commit()

    def preview(self):
        return self.R.trip_memoir_preview(TRIP)

    def docx_text(self):
        from api.services.trip_memoir_docx import build_trip_docx
        prev = self.preview()
        blob = build_trip_docx(prev)
        doc = Document(io.BytesIO(blob))
        return "\n".join(p.text for p in doc.paragraphs), doc, prev


# ══════════════════════════════════════════════════════════════════════
# NO APPROVAL GATE
# ══════════════════════════════════════════════════════════════════════
class VisibleMeansExportedTest(TimelineTestBase):

    def test_a_visible_day_exports_without_an_approval_tick(self):
        """The literal 2026-08-06 failure. include_in_memoir stays 0."""
        self.add_day(1, "2026-07-14", title="Santa Fe to Bismarck",
                     main_location="Downtown Bismarck",
                     lodging_base="Radisson Hotel on Main Street")
        text, _, _ = self.docx_text()
        self.assertIn("Santa Fe to Bismarck", text)
        self.assertIn("Downtown Bismarck", text)
        self.assertIn("Radisson Hotel on Main Street", text)

    def test_the_day_is_still_unticked_in_the_database(self):
        """Non-vacuity for the test above: it would also pass if
        something had quietly ticked the day."""
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.docx_text()
        row = self.con.execute(
            "SELECT include_in_memoir FROM trip_days WHERE id=?",
            (did,)).fetchone()
        self.assertEqual(0, row[0])

    def test_an_unticked_note_source_and_photograph_all_export(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_note(day_id=did, text="We found the headstones.")
        self.add_source(day_id=did, title="Mountrail County index")
        self.add_photo(day_id=did, narrator_caption="At the gravesite")
        text, doc, _ = self.docx_text()
        self.assertIn("We found the headstones.", text)
        self.assertIn("Mountrail County index", text)
        self.assertIn("At the gravesite", text)
        self.assertEqual(1, len(doc.inline_shapes))

    def test_every_operator_authored_day_field_is_rendered(self):
        self.add_day(1, "2026-07-14", title="Arrival",
                     main_location="Bismarck", lodging_base="Radisson",
                     morning_notes="Drove north through the grass.",
                     afternoon_notes="Found the courthouse shut.",
                     evening_notes="Walked to the river after supper.",
                     places_visited_json='["Capitol grounds", "Riverfront"]',
                     meals_json='["Walleye at the Blarney Stone"]')
        text, _, _ = self.docx_text()
        for fragment in ("Arrival", "Bismarck", "Radisson",
                         "Drove north through the grass.",
                         "Found the courthouse shut.",
                         "Walked to the river after supper.",
                         "Capitol grounds", "Riverfront",
                         "Walleye at the Blarney Stone"):
            self.assertIn(fragment, text, fragment)

    def test_a_title_only_day_is_still_printed(self):
        """`title` is rendered in the heading, not as a timeline item,
        so a day filtered on item count alone disappeared while being
        plainly visible on the timeline."""
        self.add_day(1, "2026-07-14", title="Santa Fe to Bismarck")
        text, _, prev = self.docx_text()
        self.assertEqual(0, prev["export_summary"]["day_items"])
        self.assertIn("Santa Fe to Bismarck", text)

    def test_a_day_with_only_a_number_and_a_date_is_still_printed(self):
        """A day row is visible on the timeline because it exists. Its
        number and its date ARE content, and filtering on items-or-title
        dropped a real day of the trip from the document while the
        operator could see it on screen."""
        self.add_day(2, "2026-07-15")
        text, _, prev = self.docx_text()
        self.assertEqual(0, prev["export_summary"]["day_items"])
        self.assertIn("Day 2", text)
        self.assertIn("2026-07-15", text)

    def test_every_projected_day_is_rendered_with_no_second_filter(self):
        """Six empty days, six headings. Two filters would mean two
        definitions of "a day worth printing", in two languages."""
        for idx in range(1, 7):
            self.add_day(idx, "2026-07-%02d" % (13 + idx))
        text, _, prev = self.docx_text()
        self.assertEqual(6, len(prev["part_one_timeline"]["days"]))
        for idx in range(1, 7):
            self.assertIn("Day %d" % idx, text)

    def test_six_days_stay_chronological(self):
        """Inserted out of order on purpose: the ordering must come from
        day_index/date, not from insertion."""
        for idx in (4, 1, 6, 3, 2, 5):
            self.add_day(idx, "2026-07-%02d" % (13 + idx),
                         title=f"Day{idx}Marker")
        text, _, _ = self.docx_text()
        seen = [text.index(f"Day{i}Marker") for i in range(1, 7)]
        self.assertEqual(sorted(seen), seen, "days are out of order")


# ══════════════════════════════════════════════════════════════════════
# CONVERSATIONS ARE IN, WITH BOTH SPEAKERS
# ══════════════════════════════════════════════════════════════════════
class ConversationsExportTest(TimelineTestBase):

    def test_a_day_linked_conversation_exports_once_with_both_speakers(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_conversation(day_id=did,
                              said="I walked to the river.",
                              replied="What do you remember of it?")
        text, _, prev = self.docx_text()
        self.assertIn("Christopher: I walked to the river.", text)
        self.assertIn("Lori: What do you remember of it?", text)
        self.assertEqual(1, text.count("I walked to the river."))
        self.assertEqual(1, text.count("What do you remember of it?"))

    def test_the_speaker_label_is_the_narrator_s_own_name(self):
        self.add_day(1, "2026-07-14")
        self.add_conversation(day_id="day-1")
        _, _, prev = self.docx_text()
        self.assertEqual("Christopher", prev["narrator_label"])

    def test_a_nameless_narrator_falls_back_and_does_not_guess(self):
        self.con.execute("UPDATE people SET display_name='' WHERE id=?",
                         (PERSON,))
        self.con.commit()
        self.add_day(1, "2026-07-14")
        self.add_conversation(day_id="day-1")
        text, _, prev = self.docx_text()
        self.assertEqual("Narrator", prev["narrator_label"])
        self.assertIn("Narrator: ", text)

    def test_a_system_directive_is_not_attributed_to_the_narrator(self):
        """A `[SYSTEM: ...]` turn is an instruction this system sent to
        Lori through the user channel. Printing it as the narrator's
        speech puts words in his mouth in the artefact his family reads
        -- the 2026-07-14 directive-concatenation bug, one layer out."""
        did = self.add_day(1, "2026-07-14")
        self.add_conversation(
            day_id=did,
            said="[SYSTEM: The narrator has been quiet for a while.]",
            replied="Take your time.")
        text, _, _ = self.docx_text()
        self.assertNotIn("[SYSTEM", text)
        self.assertNotIn("quiet for a while", text)
        self.assertIn("Lori: Take your time.", text)

    def test_a_rejected_placement_does_not_export(self):
        did = self.add_day(1, "2026-07-14")
        self.add_conversation(day_id=did, said="THROWN AWAY",
                              replied="ALSO THROWN AWAY",
                              status="rejected")
        text, _, _ = self.docx_text()
        self.assertNotIn("THROWN AWAY", text)


# ══════════════════════════════════════════════════════════════════════
# ONLY HIDDEN / DELETED / REJECTED / OTHER-TRIP IS EXCLUDED
# ══════════════════════════════════════════════════════════════════════
class OnlyTheFourExclusionsTest(TimelineTestBase):

    def test_hidden_material_does_not_export(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_note(day_id=did, text="HIDDEN NOTE", hidden=1)
        self.add_source(day_id=did, title="HIDDEN SOURCE", hidden=1)
        self.add_photo(day_id=did, narrator_caption="HIDDEN PHOTO",
                       hidden=1)
        text, doc, _ = self.docx_text()
        self.assertIn("Arrival", text)
        for gone in ("HIDDEN NOTE", "HIDDEN SOURCE", "HIDDEN PHOTO"):
            self.assertNotIn(gone, text, gone)
        self.assertEqual(0, len(doc.inline_shapes))

    def test_a_soft_deleted_photograph_does_not_export(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_photo(day_id=did, narrator_caption="DELETED PHOTO",
                       deleted="2026-08-01T00:00:00Z")
        text, doc, _ = self.docx_text()
        self.assertNotIn("DELETED PHOTO", text)
        self.assertEqual(0, len(doc.inline_shapes))

    def test_another_trip_s_material_does_not_export(self):
        self.con.execute(
            "INSERT INTO trips (id, person_id, title, created_at,"
            " updated_at) VALUES ('other',?, 'Other', 't','t')", (PERSON,))
        self.con.execute(
            "INSERT INTO trip_location_notes (id, trip_id, note_text,"
            " hidden, created_at, updated_at)"
            " VALUES ('n-other','other','OTHER TRIP NOTE',0,'t','t')")
        self.con.commit()
        self.add_day(1, "2026-07-14", title="Arrival")
        text, _, _ = self.docx_text()
        self.assertNotIn("OTHER TRIP NOTE", text)


# ══════════════════════════════════════════════════════════════════════
# NEEDS A DAY
# ══════════════════════════════════════════════════════════════════════
class NeedsADayTest(TimelineTestBase):

    def test_unplaced_material_appears_under_needs_a_day(self):
        self.add_day(1, "2026-07-14", title="Arrival")
        self.add_note(text="A thought with no day.")
        self.add_photo(narrator_caption="A photo with no day")
        self.add_conversation(day_id=None, said="Told with no day.",
                              replied="Go on.", status="needs_day")
        text, _, _ = self.docx_text()
        self.assertIn("Needs a day", text)
        tail = text.split("Needs a day", 1)[1]
        self.assertIn("A thought with no day.", tail)
        self.assertIn("A photo with no day", tail)
        self.assertIn("Christopher: Told with no day.", tail)

    def test_the_heading_is_absent_when_everything_has_a_day(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_note(day_id=did, text="Placed.")
        text, _, _ = self.docx_text()
        self.assertNotIn("Needs a day", text)


# ══════════════════════════════════════════════════════════════════════
# EXACTLY ONCE
# ══════════════════════════════════════════════════════════════════════
class ExactlyOnceTest(TimelineTestBase):

    def test_a_placeless_note_is_printed_once_not_twice(self):
        """It used to be collected into the preview's trip-level
        `story_notes` AND into the timeline's unplaced items, which
        printed all eleven Bismarck notes twice in one document."""
        self.add_note(text="We found the headstones.")
        text, _, prev = self.docx_text()
        self.assertEqual(1, text.count("We found the headstones."))
        self.assertEqual([], prev["story_notes"])

    def test_a_day_scoped_note_is_not_also_printed_by_the_region_walk(self):
        self.con.execute(
            "INSERT INTO trip_regions (id, trip_id, title, ord, created_at,"
            " updated_at) VALUES ('r1',?,'North Dakota',0,'t','t')", (TRIP,))
        self.con.execute(
            "INSERT INTO trip_stops (id, trip_id, trip_region_id,"
            " location_name, ord, created_at, updated_at)"
            " VALUES ('s1',?, 'r1','Stanley',0,'t','t')", (TRIP,))
        self.con.commit()
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_note(day_id=did, stop_id="s1", text="Written once.")
        text, _, _ = self.docx_text()
        self.assertEqual(1, text.count("Written once."))

    def test_a_photograph_is_embedded_once(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_photo(day_id=did, narrator_caption="Only once")
        text, doc, _ = self.docx_text()
        self.assertEqual(1, len(doc.inline_shapes))
        self.assertEqual(1, text.count("Only once"))

    def test_the_photo_appendix_is_gone(self):
        """Rule 10 over rule 11's permission: an appendix would embed
        every image a second time."""
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did)
        text, _, prev = self.docx_text()
        self.assertNotIn("Part III", text)
        self.assertNotIn("Photo Appendix", text)
        self.assertTrue(prev["part_three_photo_appendix"]["unknown"])


# ══════════════════════════════════════════════════════════════════════
# MACHINE TEXT IS LABELLED
# ══════════════════════════════════════════════════════════════════════
class MachineCaptionTest(TimelineTestBase):

    def test_a_machine_description_is_labelled_as_draft(self):
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did,
                       machine_description="A stone building with a flag.")
        text, _, prev = self.docx_text()
        item = prev["part_one_timeline"]["days"][0]["items"][0]
        self.assertEqual("machine", item["caption_source"])
        self.assertIn("Draft description (machine-written, not reviewed)",
                      text)

    def test_the_narrator_s_own_caption_is_not_labelled(self):
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did, narrator_caption="Mum and Dad, 1962",
                       machine_description="Two people outdoors.")
        text, _, prev = self.docx_text()
        item = prev["part_one_timeline"]["days"][0]["items"][0]
        self.assertEqual("narrator", item["caption_source"])
        self.assertIn("Mum and Dad, 1962", text)
        self.assertNotIn("machine-written", text)
        self.assertNotIn("Two people outdoors.", text)

    def test_an_operator_caption_wins_over_the_machine_and_is_not_labelled(self):
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did, operator_caption="The courthouse",
                       machine_description="A stone building.")
        text, _, prev = self.docx_text()
        item = prev["part_one_timeline"]["days"][0]["items"][0]
        self.assertEqual("operator", item["caption_source"])
        self.assertIn("The courthouse", text)
        self.assertNotIn("machine-written", text)


# ══════════════════════════════════════════════════════════════════════
# ONE SHARED PROJECTION
# ══════════════════════════════════════════════════════════════════════
class OneProjectionTest(TimelineTestBase):

    def test_the_document_renders_the_preview_s_own_timeline(self):
        """Rule 9, proved by mutating the preview and watching the
        document follow."""
        from api.services.trip_memoir_docx import build_trip_docx
        self.add_day(1, "2026-07-14", title="Arrival")
        prev = self.preview()
        prev["part_one_timeline"]["days"][0]["title"] = "SENTINEL"
        blob = build_trip_docx(prev)
        text = "\n".join(p.text for p in Document(io.BytesIO(blob)).paragraphs)
        self.assertIn("SENTINEL", text)
        self.assertNotIn("Arrival", text)

    def test_the_export_reads_the_same_function_the_timeline_does(self):
        """`trip_day_timeline_items` IS the visible timeline. The export
        must share it rather than keep a second interpretation of a day.
        Read from the AST, because prose in this module names the
        function too."""
        import ast
        import inspect
        from api.services import trip_repository as R
        src = inspect.getsource(R.trip_timeline_projection)
        tree = ast.parse(src.lstrip())
        called = {n.func.id for n in ast.walk(tree)
                  if isinstance(n, ast.Call) and isinstance(n.func, ast.Name)}
        self.assertIn("trip_day_timeline_items", called)
        self.assertIn("trip_day_conversation_items", called)

    def test_no_route_or_repository_writes_the_dormant_day_column(self):
        """Migration 0042 stays applied; the column must be inert. A
        writer with no reader is how a dead field comes back to life."""
        import ast
        import inspect
        import re as _re
        from pathlib import Path as _P
        from api.services import trip_repository as R

        # Scoped to the DAY lane on purpose. Notes, sources, photo links
        # and photo context all keep their own include_in_memoir
        # controls -- a file-wide scan would fire on those and would be
        # asserting something Chris did not ask for.
        self.assertNotIn(
            "include_in_memoir",
            inspect.signature(R.trip_day_update).parameters,
            "the day repository writer still accepts an approval")

        src = (_P(__file__).resolve().parent.parent
               / "server/code/api/routers/trips.py").read_text(
                   encoding="utf-8")
        tree = ast.parse(src)
        # The day patch model and the day patch handler, by name.
        for node in ast.walk(tree):
            if isinstance(node, ast.ClassDef) and node.name == "TripDayPatch":
                fields = {t.target.id for t in node.body
                          if isinstance(t, ast.AnnAssign)}
                self.assertNotIn("include_in_memoir", fields,
                                 "TripDayPatch still offers an approval")
            if isinstance(node, ast.FunctionDef) and "day" in node.name:
                body = ast.get_source_segment(src, node) or ""
                code = _re.sub(r"#.*$", "", body, flags=_re.M)
                code = _re.sub(r'"""[\s\S]*?"""', "", code)
                self.assertNotIn("include_in_memoir", code, node.name)

        # And nothing anywhere writes the column.
        for rel in ("server/code/api/routers/trips.py",
                    "server/code/api/services/trip_repository.py"):
            text = (_P(__file__).resolve().parent.parent / rel).read_text(
                encoding="utf-8")
            self.assertNotIn("UPDATE trip_days SET include_in_memoir", text)

    def test_the_dormant_column_is_not_read_on_the_export_path(self):
        """Migration 0042 stays applied so fresh installations match
        Chris's database. Nothing may depend on it."""
        import ast
        import inspect
        from api.services import trip_repository as R
        from api.services import trip_memoir_docx as D
        for mod, fn in ((R, "trip_timeline_projection"),
                        (D, "build_trip_docx")):
            src = inspect.getsource(getattr(mod, fn))
            body = ast.get_source_segment(
                src, ast.parse(src.lstrip()).body[0]) or src
            # Comment- and docstring-stripped: the retirement notes name
            # the column to explain that it is dormant, and a raw scan
            # would fire on the explanation.
            import re
            code = re.sub(r"#.*$", "", body, flags=re.M)
            code = re.sub(r'"""[\s\S]*?"""', "", code)
            self.assertNotIn("include_in_memoir", code, fn)

    def test_the_default_projection_carries_no_filesystem_path(self):
        """`_day_photo_items` feeds the LIVE day-timeline endpoint too,
        so a path put there for the document's benefit leaks to the
        operator interface -- which is what `test_trip_placement`'s
        boundary guard caught."""
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did)
        self.add_photo()
        proj = self.R.trip_timeline_projection(TRIP)
        for group in list(proj["days"]) + [proj["unplaced"]]:
            for item in group["items"]:
                self.assertNotIn("image_path", item)

    def test_the_builder_copy_does_carry_one(self):
        """Non-vacuity for the test above."""
        did = self.add_day(1, "2026-07-14")
        self.add_photo(day_id=did)
        proj = self.R.trip_timeline_projection(TRIP, with_image_paths=True)
        self.assertTrue(proj["days"][0]["items"][0]["image_path"])


class CountsTest(TimelineTestBase):

    def test_the_summary_reports_plain_counts_not_approvals(self):
        did = self.add_day(1, "2026-07-14", title="Arrival")
        self.add_day(2, "2026-07-15")
        self.add_note(day_id=did, text="one")
        self.add_source(day_id=did, title="two")
        self.add_photo(day_id=did)
        s = self.preview()["export_summary"]
        self.assertEqual(2, s["days"])
        self.assertEqual(1, s["notes"])
        self.assertEqual(1, s["sources"])
        self.assertEqual(1, s["photos"])
        # note + source + photo. The day TITLE is not an item -- it is
        # the heading -- so a title-only day counts zero items and is
        # still printed. `test_a_title_only_day_is_still_printed` is the
        # guard for that.
        self.assertEqual(3, s["day_items"])
        for gone in ("days_in", "days_out", "notes_in", "sources_in",
                     "photos_in", "days_approved_empty",
                     # Renamed: these required a memoir tick as well,
                     # which decides nothing about this document.
                     "notes_hidden_approved", "sources_hidden_approved",
                     "photos_hidden_approved"):
            self.assertNotIn(gone, s, gone)
        for kept in ("notes_hidden", "sources_hidden", "photos_hidden"):
            self.assertIn(kept, s, kept)


if __name__ == "__main__":
    unittest.main(verbosity=2)
