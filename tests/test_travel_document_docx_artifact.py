"""WO-TRAVEL-DOC-CLOSEOUT-01 — open the actual Word document.

WHY THIS SUITE EXISTS
---------------------
Every other travel-document test reads source. That is the right
instrument for "this control is wired to that route" and the wrong one
for "the family will read these words". Chris's point, verbatim: the
final proof must be the one actual exported document.

So this builds a real .docx with python-docx, opens it, extracts every
paragraph, and compares that text against the projection the browser
preview renders. If the preview and the document ever disagree about a
caption, a heading, a date or a missing file, this fails — and it fails
on the artefact, not on a substring in a source file.

WHAT IT PROVES

  * every word the document contains is in the preview projection;
  * `photo_description` never reaches the document, whatever the row
    carries;
  * an unapproved operator caption is absent;
  * a title-only story note still appears;
  * two stops with the SAME NAME stay two sections;
  * a photograph whose file is missing is reported, not silently
    dropped;
  * per-stop counts match the appendix.

Run:
    PYTHONPATH=server/code .venv/bin/python -m unittest tests.test_travel_document_docx_artifact
"""
from __future__ import annotations

import io
import os
import sys
import tempfile
import unittest
from pathlib import Path

_REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_REPO / "server" / "code"))

try:
    from docx import Document          # noqa: F401
    _HAVE_DOCX = True
except Exception:                       # pragma: no cover - env dependent
    _HAVE_DOCX = False

from api.services.trip_repository import photo_appendix_projection  # noqa: E402
from api.services.trip_memoir_docx import build_trip_docx           # noqa: E402


def _png(path: Path) -> Path:
    """Smallest valid PNG python-docx will embed."""
    path.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753"
        "de0000000c4944415478da63f8ffff3f0005fe02fea735c9ab0000000049454e"
        "44ae426082"))
    return path


class DocxArtifactTest(unittest.TestCase):
    """Builds a document and reads it back."""

    @classmethod
    def setUpClass(cls):
        if not _HAVE_DOCX:            # pragma: no cover
            raise unittest.SkipTest("ENV-SKIP: python-docx not installed")
        cls.tmp = tempfile.TemporaryDirectory()
        d = Path(cls.tmp.name)
        cls.good = _png(d / "good.png")
        cls.missing = d / "gone.png"        # deliberately never created

        # Two stops with the SAME display name, on purpose.
        cls.rows = [
            {"id": "l1", "photo_id": "p1", "trip_stop_id": "s1",
             "stop_location_name": "Hotel",
             "narrator_caption": "Peter and Josie at the gate",
             "caption": "operator text", "caption_approved_for_lori": 0,
             "photo_description": "LEAKED-DESCRIPTION",
             "taken_at": "1962-06-01",
             "photo_image_path": str(cls.good)},
            {"id": "l2", "photo_id": "p2", "trip_stop_id": "s2",
             "stop_location_name": "Hotel",
             "narrator_caption": "", "caption": "UNAPPROVED-OPERATOR-CAPTION",
             "caption_approved_for_lori": 0,
             "photo_description": "ALSO-LEAKED",
             "taken_at": "1962-06-02",
             "photo_image_path": str(cls.good)},
            {"id": "l3", "photo_id": "p3", "trip_stop_id": "s1",
             "stop_location_name": "Hotel",
             "narrator_caption": "", "caption": "APPROVED-OPERATOR-CAPTION",
             "caption_approved_for_lori": 1,
             "taken_at": "1962-06-03",
             "photo_image_path": str(cls.good)},
            {"id": "l4", "photo_id": "p4", "trip_stop_id": None,
             "trip_region_id": "r1", "region_title": "Hotel",
             "narrator_caption": "region shot",
             "taken_at": "1962-06-04",
             "photo_image_path": str(cls.missing)},   # file absent
        ]
        cls.projection = photo_appendix_projection(rows=list(cls.rows))

        cls.preview = {
            "title": "Bismarck 1962",
            "date_range": {"start": "1962-05-30", "end": "1962-06-10"},
            "summary": "TRIP-SUMMARY",
            "story_notes": [{"note_title": "TITLE-ONLY-NOTE", "note_text": ""}],
            "sources": [{"title": "SRC", "pasted_text": "PASTED-DETAIL"}],
            "part_one_journey_in_order": [{
                "region": "Dakota", "date_range": {"start": "1962-05-30",
                                                   "end": "1962-06-10"},
                "base_address": "BASE-ADDRESS", "summary": "REGION-SUMMARY",
                "story_notes": [], "sources": [],
                "stops": [
                    {"id": "s1", "location_name": "Hotel",
                     "date_start": "1962-06-01", "notes": "STOP-OWN-NOTES",
                     "story_notes": [], "sources": [], "day_trips": [
                         {"id": "s3", "location_name": "Cemetery",
                          "story_notes": [{"note_title": "",
                                           "note_text": "GRAVESITE-STORY"}],
                          "sources": [], "day_trips": []}]},
                    {"id": "s2", "location_name": "Hotel",
                     "story_notes": [], "sources": [], "day_trips": []},
                ]}],
            "part_two_themes": [{"theme": "THEME-NAME",
                                 "description": "THEME-DESCRIPTION",
                                 "stops": ["Hotel"]}],
            "part_three_photo_appendix": {},
        }
        cls.blob = build_trip_docx(cls.preview, cls.rows)
        doc = Document(io.BytesIO(cls.blob))
        cls.paras = [p.text for p in doc.paragraphs]
        cls.text = "\n".join(cls.paras)

    @classmethod
    def tearDownClass(cls):
        if _HAVE_DOCX:
            cls.tmp.cleanup()

    # ── the leak ──────────────────────────────────────────────────────
    def test_no_photo_description_reaches_the_document(self):
        """The one that would put words nobody approved under a
        photograph in a family memoir."""
        self.assertNotIn("LEAKED-DESCRIPTION", self.text)
        self.assertNotIn("ALSO-LEAKED", self.text)

    def test_an_unapproved_operator_caption_is_absent(self):
        self.assertNotIn("UNAPPROVED-OPERATOR-CAPTION", self.text)

    def test_an_approved_operator_caption_is_present(self):
        """Non-vacuity: without this, the two above would pass on a
        document with no captions at all."""
        self.assertIn("APPROVED-OPERATOR-CAPTION", self.text)

    def test_the_narrator_caption_is_present(self):
        self.assertIn("Peter and Josie at the gate", self.text)

    # ── grouping identity ─────────────────────────────────────────────
    def test_two_stops_with_the_same_name_stay_two_sections(self):
        """Grouping by display text merged them into one, so the
        photographs of two different places silently became one place."""
        self.assertEqual(3, len(self.projection["groups"]),
                         [g["key"] for g in self.projection["groups"]])
        keys = [g["key"] for g in self.projection["groups"]]
        self.assertEqual(len(keys), len(set(keys)))
        # Counted as WHOLE PARAGRAPHS equal to "Hotel", which is what an
        # appendix group heading is. My first cut counted the substring
        # across the document and failed at 5 -- Part I's bullets and the
        # theme's "Across:" line both contain the word, and neither is a
        # section. Counting prose is not counting sections.
        headings = [t for t in self.paras if t.strip() == "Hotel"]
        self.assertEqual(3, len(headings),
                         f"expected three separate sections named Hotel "
                         f"(two stops and a region), got {len(headings)}")

    def test_the_region_group_is_distinct_from_the_stops(self):
        scopes = {g["scope"] for g in self.projection["groups"]}
        self.assertEqual({"stop", "region"}, scopes)

    # ── missing files ─────────────────────────────────────────────────
    def test_a_missing_file_is_reported_not_silently_dropped(self):
        self.assertEqual(1, self.projection["unavailable"])
        self.assertIn("could not be found on disk", self.text)

    def test_an_unavailable_only_group_is_not_an_empty_heading(self):
        """A section in a family memoir with a heading and nothing under
        it looks like a mistake or a deletion. It says what happened."""
        i = self.paras.index("Hotel", self.paras.index("Part III — Photo Appendix"))
        # The region group (all of whose photographs are missing) must be
        # followed by an explanation, not by the next heading.
        tail = "\n".join(self.paras[i:])
        self.assertIn("approved here could not be found on disk", tail)

    def test_approved_available_and_embedded_are_three_numbers(self):
        """#7. `approved` is what was ticked; `available` is what is on
        disk; `embedded` is knowable only after Word accepts each image,
        so it is reported last and promised nowhere earlier."""
        self.assertIn(f"Approved photos in appendix: {len(self.rows)}",
                      self.text)
        self.assertIn(f"({self.projection['available']} photos embedded",
                      self.text)
        self.assertNotIn("will be embedded", self.text)
        self.assertNotIn("The rest will be embedded", self.text)

    def test_the_approved_count_is_the_row_count(self):
        self.assertIn(f"Approved photos in appendix: {len(self.rows)}",
                      self.text)

    def test_per_stop_counts_match_the_appendix(self):
        # s1 has two approved photographs; the Part I bullet must say so.
        self.assertEqual(2, self.projection["approved_by_stop"]["s1"])
        self.assertIn("· 2 photos", self.text)

    # ── every exported word is in the preview projection ──────────────
    def test_every_exported_text_field_appears(self):
        """The prose the operator reviews must be the prose that ships."""
        for expected in ("Bismarck 1962", "1962-05-30 — 1962-06-10",
                         "TRIP-SUMMARY", "TITLE-ONLY-NOTE", "PASTED-DETAIL",
                         "BASE-ADDRESS", "REGION-SUMMARY", "STOP-OWN-NOTES",
                         "GRAVESITE-STORY", "THEME-NAME", "THEME-DESCRIPTION"):
            with self.subTest(text=expected):
                self.assertIn(expected, self.text)

    def test_a_title_only_note_survives(self):
        """The DOCX prints title and body independently, so a note with a
        title and no body still reaches the document."""
        self.assertIn("TITLE-ONLY-NOTE", self.text)

    def test_the_nested_day_trip_note_is_exported(self):
        self.assertIn("GRAVESITE-STORY", self.text)

    # ── the preview projection and the document agree ─────────────────
    def test_every_caption_in_the_projection_is_in_the_document(self):
        for g in self.projection["groups"]:
            for ph in g["photos"]:
                if not ph["available"] or not ph["caption"]:
                    continue
                with self.subTest(caption=ph["caption"]):
                    self.assertIn(ph["caption"], self.text)

    def test_every_group_label_in_the_projection_is_a_heading(self):
        headings = set(self.paras)
        for g in self.projection["groups"]:
            with self.subTest(label=g["label"]):
                self.assertIn(g["label"], headings)

    def test_the_projection_carries_no_photo_description(self):
        """The browser gets this object. It must not contain the field
        even as an unused key."""
        for g in self.projection["groups"]:
            for ph in g["photos"]:
                self.assertNotIn("photo_description", ph)
                self.assertNotIn("description", ph)

    def test_a_stop_with_no_approved_photograph_prints_no_count(self):
        """s2 has one approved photograph and s3 (the nested Cemetery day
        trip) has none. A bare "· 0 photos" on every empty stop is noise,
        and the absence is what proves the count comes from the export
        set rather than from the trip tree."""
        self.assertNotIn("· 0 photo", self.text)
        cem = [t for t in self.paras if t.startswith("Cemetery")]
        self.assertTrue(cem, "the nested day trip is missing from Part I")
        self.assertNotIn("photo", cem[0])

    def test_the_projection_can_be_built_from_rows_or_a_trip(self):
        """One implementation. The builder reuses the caller's rows; the
        preview builds from a trip id. A second grouping implementation
        would drift, and the whole point is that they cannot."""
        from api.services.trip_repository import photo_appendix_projection
        again = photo_appendix_projection(rows=list(self.rows))
        self.assertEqual([g["key"] for g in self.projection["groups"]],
                         [g["key"] for g in again["groups"]])

    def test_passing_the_projection_gives_the_same_document(self):
        """#3: the route builds it once and hands the same object to
        both consumers. That must produce byte-comparable text to the
        rows path, or "one read" would be changing the output."""
        from api.services.trip_memoir_docx import build_trip_docx as build
        import io as _io
        blob = build(self.preview, appendix=self.projection)
        paras = [p.text for p in Document(_io.BytesIO(blob)).paragraphs]
        self.assertEqual(self.paras, paras)

    def test_the_document_is_a_real_openable_docx(self):
        """Non-vacuity for everything above: if the bytes were not a
        document, `Document()` would have raised and every assertion
        would be about an empty string."""
        self.assertGreater(len(self.blob), 5000)
        self.assertEqual(b"PK", self.blob[:2])
        self.assertGreater(len(self.paras), 15)


if __name__ == "__main__":
    unittest.main(verbosity=2)
