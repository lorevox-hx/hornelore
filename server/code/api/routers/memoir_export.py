"""
Lorevox Memoir Export Router
=============================
Provides server-side DOCX export for memoir content.

Endpoints:
  POST  /api/memoir/export-docx  — accept memoir JSON, return .docx file

Design rules:
  - Export reflects exactly what the user sees (threads or draft).
  - Scaffold placeholder content is never exported.
  - Meaning sections (Turning Points, Hard Moments, etc.) become DOCX headings.
  - Structural sections (Family & Relationships, Work, etc.) become secondary headings.
  - Draft state is rendered as plain prose paragraphs with section headers.
  - Threads state is rendered as grouped bullet lists per section.
  - Media Builder: attached_photos inlines images after section headings (graceful skip on error).

WO-ML-04 / Phase 4B (2026-05-07) — bilingual memoir export:
  - target_language="en" (default): English-only output, byte-stable
    with pre-Phase-4B callers.
  - target_language="es": Spanish-only output. Each section's items +
    prose pass through services.translation.translate_text() before
    rendering. Title + subtitle render with their Spanish equivalents.
    On translation failure, the affected section falls back to its
    English source text so the caller always gets a usable docx.
  - target_language="bilingual": both languages in the same document.
    Rendered as paired blocks per section (English first, Spanish
    immediately below) so a bilingual reader can follow either side.
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from .. import flags

logger = logging.getLogger("memoir_export")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

router = APIRouter(prefix="/api/memoir", tags=["memoir-export"])


# ── Flag gate ─────────────────────────────────────────────────────────────────

def _require_enabled() -> None:
    """WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01: raise 404 when
    HORNELORE_MEMOIR_EXPORT_ENABLED is off. Mirrors the trips/photos
    posture — a disabled surface does not advertise itself."""
    if not flags.memoir_export_enabled():
        raise HTTPException(status_code=404, detail="Not found")


# ── Request models ─────────────────────────────────────────────────────────────

class MemoirSection(BaseModel):
    """A single named section with zero or more thread items.

    `sources` (2026-08-19) carries one opaque provenance digest per item,
    parallel to `items`. It is set only by the server's captured-story
    harvest, is never rendered into the visible document, and never
    contains a narrator id -- see `_story_source_digest`. Client sections
    simply leave it empty.
    """
    id: str
    label: str
    items: List[str] = Field(default_factory=list)
    sources: List[str] = Field(default_factory=list)
    #: Per-item ISO-639-1 source language, parallel to `items`.
    #: 2026-08-19: the pipeline assumed every source item was English,
    #: so a Spanish story was handed to the translator as English and a
    #: Spanish-to-Spanish "translation" was reported as success. Set by
    #: the server harvests from the stored language; empty on client
    #: sections, which fall back to `req.source_language` as before.
    languages: List[str] = Field(default_factory=list)


class AttachedPhoto(BaseModel):
    """A photo attached to a memoir section (Media Builder — Task 4).

    WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5.2): the
    client no longer holds file-path authority. ``file_path`` is kept
    ONLY for wire-compat with old clients that still send it — the
    server never reads it (not for logging, validation, or rendering).
    The on-disk location is resolved server-side from ``media_id``
    through the media table (see _resolve_media_photo_path)."""
    media_id: str
    section_key: str
    file_path: Optional[str] = None   # IGNORED — wire-compat only, never read
    description: str = ""
    taken_at: str = ""


class MemoirExportRequest(BaseModel):
    """
    Shape sent by the frontend's memoirExportDOCX() function.
    memoir_state: "threads" | "draft"
    narrator_name: display name for the document title
    sections: populated sections only (empty sections are pre-filtered by the caller)
    prose: flat prose string for draft state (paragraphs joined by \\n\\n)
    arc_roles: which narrative arc parts are present (display only, optional)
    attached_photos: photos to inline at their section (empty list = no change in behavior)

    WO-ML-04 / Phase 4B (2026-05-07) — bilingual memoir export:
      source_language: ISO-639-1 code of the source content. Default
        "en". Used as the source for translation calls; also informs
        the chrome (subtitle wording) when target_language matches.
      target_language: ISO-639-1 code OR "bilingual". Default "en".
        - "en" → no translation, English-only output (byte-stable
          with pre-Phase-4B callers).
        - "es" → translate every section item + prose to Spanish
          via services.translation.translate_text() before render.
        - "bilingual" → render English + Spanish side-by-side.
    """
    narrator_name: str = Field(default="Narrator")
    memoir_state: str = Field(default="threads")
    # WO-MEMOIR-STORY-CANDIDATES-WIRE-01 (2026-07-06): when person_id
    # is present, the server harvests operator-cleared captured
    # stories (story_candidates with review_status promoted or
    # memoir_only) and appends them as era-grouped sections in the
    # narrator's OWN words. Absent person_id = byte-stable with every
    # pre-wire caller. include_captured_stories=False opts out.
    person_id: Optional[str] = Field(default=None)
    include_captured_stories: bool = Field(default=True)
    # WO-MEMOIR-TRIP-STORY-LANE-01 (2026-07-27): approved Travel Doc
    # trip stories (trip_location_notes.include_in_memoir=1) join the
    # narrator memoir as their own clearly-sourced sections. This is a
    # DB read. It NEVER writes a travel_doc_modal turn into the
    # life-story archive -- the two-surface rule of 2026-07-09 holds
    # (see tests/test_modal_archive_boundary.py). Opt out per request.
    include_trip_stories: bool = Field(default=True)
    sections: List[MemoirSection] = Field(default_factory=list)
    prose: Optional[str] = Field(default=None)
    arc_roles: List[str] = Field(default_factory=list)
    attached_photos: List[AttachedPhoto] = Field(default_factory=list)
    source_language: str = Field(default="en")
    target_language: str = Field(default="en")


# ── Helpers ────────────────────────────────────────────────────────────────────

# Colour constants for the Lorevox brand tone (dark warm palette)
# Guarded: RGBColor only exists when python-docx is installed.
if _DOCX_AVAILABLE:
    _DARK_BROWN = RGBColor(0x3B, 0x2A, 0x1A)   # heading primary
    _WARM_GREY  = RGBColor(0x5A, 0x55, 0x50)   # heading secondary
    _GOLD       = RGBColor(0xAA, 0x88, 0x44)   # accent line / arc label
else:
    _DARK_BROWN = _WARM_GREY = _GOLD = None


def _photos_for_section(req: MemoirExportRequest, section_key: str) -> List[AttachedPhoto]:
    """Return all photos attached to a given memoir section key."""
    return [p for p in req.attached_photos if p.section_key == section_key]


# ── WO-ML-04 / Phase 4B helpers ───────────────────────────────────────────────

# Title + subtitle translations for the docx chrome. The narrator's
# memoir CONTENT is translated by the LLM service; this dict handles
# the deterministic boilerplate so we don't burn LLM tokens on
# four-word strings that are stable across every memoir.
_CHROME_STRINGS: Dict[str, Dict[str, str]] = {
    "en": {
        "threads_title":    "Story Threads — {narrator}",
        "threads_subtitle": "Building Blocks Collected",
        "draft_title":      "Memoir Draft — {narrator}",
        "draft_subtitle":   "Your Words, Taking Shape",
        "story_arc_label":  "Story arc: {roles}",
        "photos_heading":   "Photos",
    },
    "es": {
        "threads_title":    "Hilos de Historia — {narrator}",
        "threads_subtitle": "Piezas Recogidas",
        "draft_title":      "Borrador de Memorias — {narrator}",
        "draft_subtitle":   "Tus Palabras, Tomando Forma",
        "story_arc_label":  "Arco narrativo: {roles}",
        "photos_heading":   "Fotos",
    },
}


def _chrome(target_lang: str, key: str) -> str:
    """Return a chrome string in the requested language, falling back
    to English when the target language has no entry. Stays
    byte-stable when target_lang == 'en'."""
    table = _CHROME_STRINGS.get(target_lang) or _CHROME_STRINGS["en"]
    if key in table:
        return table[key]
    return _CHROME_STRINGS["en"].get(key, "")


def _normalize_target_lang(req: MemoirExportRequest) -> str:
    """Return one of: 'en' (default, byte-stable) | 'es' | 'bilingual'.
    Coerces unknown values to 'en' so a malformed request still
    produces an English memoir rather than crashing."""
    raw = (req.target_language or "en").strip().lower()
    if raw in ("en", "es", "bilingual"):
        return raw
    logger.warning(
        "[memoir-docx][lang] unsupported target_language=%r — defaulting to 'en'",
        req.target_language,
    )
    return "en"


def _translate_request_content(
    req: MemoirExportRequest,
    target_lang: str,
) -> MemoirExportRequest:
    """Build a new MemoirExportRequest with all narrator content
    translated to `target_lang`. Section labels, item bullets, and
    prose paragraphs all pass through services.translation. Photo
    metadata and arc_roles are NOT translated (stable identifiers /
    operator-side labels).

    On translation failure for any single field, that field falls
    back to its English source — so the docx always renders.

    Used for target_lang='es' export. NOT used for 'bilingual' (that
    path renders both source + translated content side-by-side and
    handles its own translation calls inline).
    """
    # Lazy import — translation service hits the LLM and pulls in
    # urllib network paths; keep memoir export importable in
    # contexts where the LLM stack isn't available (eg unit tests).
    from ..services import translation as _translation

    source_lang = (req.source_language or "en").strip().lower() or "en"

    def _t(text: Optional[str], src: Optional[str] = None) -> str:
        if not text:
            return text or ""
        try:
            return _translation.translate_text(
                text,
                source_lang=(src or source_lang),
                target_lang=target_lang,
                narrator_name=req.narrator_name or None,
            )
        except Exception as exc:
            logger.warning(
                "[memoir-docx][translate] failed text_len=%d err=%s — passing through",
                len(text), exc,
            )
            return text

    translated_sections: List[MemoirSection] = []
    for sec in req.sections:
        # ── PER-ITEM SOURCE LANGUAGE, 2026-08-19 ────────────────────
        #
        # Every item used to be translated with the request-level
        # `source_language`, so a Spanish story in an otherwise English
        # memoir was submitted to the translator AS ENGLISH -- and an
        # item already in the target language was translated to itself,
        # burning a call to produce the text it started with.
        #
        # `sec.languages[i]` is the item's own language where the server
        # recorded one; client sections have none and keep the previous
        # request-level behaviour exactly.
        _langs = list(sec.languages or [])

        def _t_item(text: str, idx: int) -> str:
            item_lang = (_langs[idx] if idx < len(_langs) else "") or source_lang
            if (item_lang or "").strip().lower() == target_lang:
                return text or ""      # already in the requested language
            return _t(text, src=item_lang)

        translated_sections.append(MemoirSection(
            id=sec.id,
            label=_t(sec.label),
            items=[_t_item(item, i) for i, item in enumerate(sec.items or [])],
            # Provenance survives translation. Dropping it here meant a
            # Spanish-only export lost every source digest -- the same
            # memoir, the same reviewed stories, and no way to trace a
            # paragraph back. Digests are opaque ids, not prose, so they
            # are carried across untouched. 2026-08-19.
            sources=list(sec.sources or []),
            languages=list(sec.languages or []),
        ))

    translated_prose: Optional[str] = None
    if req.prose:
        # Translate paragraph-by-paragraph so the cache hits on
        # individual paragraphs (a memoir reuses paragraphs across
        # exports more often than the whole prose blob).
        paragraphs = req.prose.split("\n\n")
        out_paragraphs = [_t(p) if p.strip() else p for p in paragraphs]
        translated_prose = "\n\n".join(out_paragraphs)

    return MemoirExportRequest(
        narrator_name=req.narrator_name,
        memoir_state=req.memoir_state,
        sections=translated_sections,
        prose=translated_prose,
        arc_roles=req.arc_roles,
        attached_photos=req.attached_photos,
        source_language=req.source_language,
        target_language=target_lang,
    )


# ── WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5.3) ──────────────
# Server-side media authority. The client supplies only media_id; the
# on-disk path comes from the media table (db.get_media_item — the same
# authority /api/media/upload writes and /api/media/file/{id} serves
# from) and must stay contained within the configured media root.

# Mirrors _ALLOWED_MIME_PREFIXES in routers/media.py (the upload-side
# allowlist). Kept local so importing this module never drags in the
# media router's FastAPI machinery.
_IMAGE_MIME_PREFIXES = (
    "image/jpeg", "image/png", "image/webp", "image/heic",
    "image/heif", "image/gif", "image/bmp", "image/tiff",
)


def _media_root() -> Path:
    """Resolve the configured media root: MEDIA_DIR env when set, else
    the established DATA_DIR/media fallback (the same location
    routers/media.py stores uploads under). Read at call time so tests
    can point it at a tempdir via the environment."""
    raw = (os.environ.get("MEDIA_DIR") or "").strip()
    if raw:
        return Path(raw).expanduser().resolve()
    data_dir = Path(os.environ.get("DATA_DIR", "data")).expanduser()
    return (data_dir / "media").resolve()


def _resolve_media_photo_path(photo: AttachedPhoto, person_id: Optional[str]) -> Path:
    """Resolve an attached photo's on-disk path SERVER-SIDE from its
    media_id. The client-supplied file_path is never consulted.

    Containment contract (fail-loud 422, never silent skip):
      - media row must exist;
      - when the request carries person_id, the row must belong to it;
      - stored MIME must be image-compatible;
      - relative stored filenames join to the media root; absolute ones
        are accepted only when they resolve inside the media root;
      - Path.resolve(strict=True) — so symlinks are flattened and a
        symlink escaping the root is rejected by the containment check;
      - must be a regular file.
    Only the path returned here may reach doc.add_picture()."""
    from .. import db as _db

    item = _db.get_media_item(photo.media_id)
    if item is None:
        raise HTTPException(
            status_code=422,
            detail=f"unknown media_id '{photo.media_id}' — not in media table",
        )

    if person_id and item.get("person_id") != person_id:
        raise HTTPException(
            status_code=422,
            detail=(
                f"media '{photo.media_id}' does not belong to person "
                f"'{person_id}'"
            ),
        )

    mime = (item.get("mime") or "").strip().lower()
    if not any(mime.startswith(p) for p in _IMAGE_MIME_PREFIXES):
        raise HTTPException(
            status_code=422,
            detail=f"media '{photo.media_id}' has non-image mime '{mime}'",
        )

    root = _media_root()
    stored = (item.get("filename") or "").strip()
    if not stored:
        raise HTTPException(
            status_code=422,
            detail=f"media '{photo.media_id}' has no stored filename",
        )

    candidate = Path(stored)
    if not candidate.is_absolute():
        candidate = root / candidate

    try:
        resolved = candidate.resolve(strict=True)
    except (FileNotFoundError, OSError) as exc:
        raise HTTPException(
            status_code=422,
            detail=f"media '{photo.media_id}' file missing on disk: {exc}",
        )

    try:
        resolved.relative_to(root)
    except ValueError:
        raise HTTPException(
            status_code=422,
            detail=(
                f"media '{photo.media_id}' resolves outside the media root"
            ),
        )

    if not resolved.is_file():
        raise HTTPException(
            status_code=422,
            detail=f"media '{photo.media_id}' is not a regular file",
        )

    return resolved


def _resolve_attached_photos(req: MemoirExportRequest) -> Dict[str, Path]:
    """Resolve every attached photo up-front, before any rendering.
    Returns media_id → contained server-resolved path. Raises 422 on
    the first authority/containment failure."""
    resolved: Dict[str, Path] = {}
    for photo in req.attached_photos:
        resolved[photo.media_id] = _resolve_media_photo_path(photo, req.person_id)
    return resolved


def _add_photo_to_doc(doc: Any, photo: AttachedPhoto, resolved: Optional[Path]) -> None:
    """
    Insert photo inline in the document.

    WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01: ``resolved`` is the
    server-resolved, root-contained path from _resolve_media_photo_path
    — the ONLY path that ever reaches doc.add_picture(). photo.file_path
    is never read. A correctly-authorized but corrupt/unsupported image
    still skips gracefully with a warning (never reads another path).
    """
    if resolved is None:
        # Defensive: photo without a resolved entry never renders.
        logger.warning(
            "[memoir-docx] no resolved path for media %s — skipping",
            photo.media_id,
        )
        return
    try:
        doc.add_picture(str(resolved), width=Inches(3.5))
        # Caption paragraph
        caption_parts = []
        if photo.description:
            caption_parts.append(photo.description)
        if photo.taken_at:
            caption_parts.append(photo.taken_at)
        if caption_parts:
            cap = doc.add_paragraph(" — ".join(caption_parts))
            if cap.runs:
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.italic = True
                cap.runs[0].font.color.rgb = _WARM_GREY
    except Exception as exc:
        logger.warning("[memoir-docx] Could not add photo %s: %s — skipping", photo.media_id, exc)


# ── DOCX builders ──────────────────────────────────────────────────────────────

def _assert_translation_covered(req: "MemoirExportRequest",
                                translated: "MemoirExportRequest",
                                target_lang: str) -> None:
    """Refuse to present untranslated evidence as a translated memoir.

    WO-LORI-CONVERSATION-TO-LIFE-MAP-MEMOIR-01 (2026-08-19).

    `_translate_request_content` falls back to the SOURCE text whenever
    the translation service fails, which keeps the export from crashing
    and is right on its own terms. What was wrong is that the result was
    then delivered as a successful Spanish memoir: an operator asks for
    Spanish, receives a document, and has no way to tell that the
    narrator's own words in it are still English.

    Only SERVER EVIDENCE is checked. Operator-authored prose is the
    operator's to write in whatever language they chose, and an item
    already IN the target language is correctly returned unchanged --
    which is why per-item `languages` had to be carried this far.

    Raises 503 when requested evidence came back untranslated.
    """
    tgt = (target_lang or "").strip().lower()
    if tgt not in ("es",):        # "en" translates nothing; bilingual keeps both
        return
    by_id = {s.id: s for s in (translated.sections or [])}
    stale: List[str] = []
    for sec in _server_evidence_sections_of(req):
        out = by_id.get(sec.id)
        out_items = list(getattr(out, "items", None) or [])
        langs = list(getattr(sec, "languages", None) or [])
        for idx, item in enumerate(sec.items or []):
            item_lang = (langs[idx] if idx < len(langs) else "") or "en"
            if item_lang == tgt:
                continue          # already in the requested language
            rendered = out_items[idx] if idx < len(out_items) else ""
            if not rendered or rendered == item:
                stale.append(f"{sec.id}:{idx}")
    if stale:
        logger.error(
            "[memoir-docx] translation incomplete for %d evidence item(s): %s",
            len(stale), ",".join(stale[:8]))
        raise HTTPException(
            status_code=503,
            detail=("the reviewed evidence could not be translated to %s — "
                    "export refused rather than deliver an untranslated "
                    "document as a translated one" % tgt))


def _stamp_source_provenance(doc, req: "MemoirExportRequest") -> None:
    """Record which captured stories this document was built from.

    WO-LORI-CONVERSATION-TO-LIFE-MAP-MEMOIR-01 Commit 3 (2026-08-19).

    The audit found exported memoirs completely flat: once written, no
    paragraph could be traced back to the story, turn or conversation it
    came from, so nothing downstream could detect a duplicate or repair a
    mistake. Provenance now travels with the artifact.

    It is written to the document's `comments` core property -- metadata,
    not a page -- so a family reading the memoir never sees it. The values
    are opaque digests, never narrator ids; an operator matches a section
    to its source by digesting the candidate id again.

    Never raises: a memoir must not fail to export over its own metadata.
    """
    try:
        # POSITIONAL, not a bare list. An unordered set of digests proves
        # only that sources were used; it cannot say WHICH paragraph came
        # from which candidate, which is the question an operator asks
        # when something in the document looks wrong. The mapping is
        # `section-id:item-index=digest`.
        pairs = []
        for sec in (req.sections or []):
            sources = getattr(sec, "sources", None) or []
            if not sources:
                continue
            sec_id = str(getattr(sec, "id", "") or "?")
            for idx, digest in enumerate(sources):
                if digest:
                    pairs.append(f"{sec_id}:{idx}={digest}")
        if not pairs:
            return
        doc.core_properties.comments = (
            "lorevox-story-sources: " + ";".join(pairs))
    except Exception as exc:
        # NOT best-effort any more, 2026-08-19. Silently shipping a
        # document whose reviewed evidence carries no provenance produces
        # exactly the artifact this lane was built to stop being: prose
        # that cannot be traced to the story it came from. If evidence is
        # present and cannot be stamped, the export fails loudly.
        logger.error("[memoir-docx] provenance stamp FAILED: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="reviewed evidence could not be stamped with its "
                   "provenance — export refused")


def _server_evidence_sections_of(req: "MemoirExportRequest") -> List["MemoirSection"]:
    """The server-harvested evidence sections in this request.

    Both lanes: reviewed captured stories and operator-approved Travel
    Document stories. Identified by the reserved id namespaces, which
    only the server may use -- client sections wearing either are
    stripped before the harvests are appended, so anything left here is
    genuinely server-authored.

    Renamed from `_captured_story_sections_of` on 2026-08-19: it matched
    `captured_stories*` only, so a draft export lost every approved trip
    story as well.
    """
    return [s for s in (req.sections or []) if _is_server_evidence_section(s)]


def _render_evidence_into_draft(doc, req: "MemoirExportRequest") -> None:
    """Render server-owned evidence into a DRAFT-state document.

    WO-LORI-CONVERSATION-TO-LIFE-MAP-MEMOIR-01 (2026-08-19).

    The draft builders render `req.prose` and ignore `req.sections`
    entirely -- correct while sections were purely the threads view of
    the same content, and wrong once the server began APPENDING evidence
    as sections. A narrator exporting in draft state got a memoir with
    every reviewed story and every approved trip story silently missing,
    while the same narrator in threads state got all of them. Same data,
    same review, two different documents.

    Rendered AFTER the operator's prose, under their own headings, so the
    authored narrative keeps its shape and the sourced evidence follows
    it rather than being woven in.
    """
    for sec in _server_evidence_sections_of(req):
        h = doc.add_heading(sec.label, level=1)
        try:
            h.runs[0].font.color.rgb = _DARK_BROWN
        except Exception:
            pass
        for item in (sec.items or []):
            doc.add_paragraph(item)
        doc.add_paragraph()  # spacer


def _render_evidence_into_bilingual_draft(doc, req: "MemoirExportRequest",
                                          translated: "MemoirExportRequest") -> None:
    """Render server-owned evidence ONCE PER LANGUAGE in a bilingual draft.

    The first cut called the monolingual renderer here, so a bilingual
    document showed the narrator's own words in the source language only
    -- while every surrounding paragraph appeared in both. A bilingual
    memoir exists so a Spanish-reading grandchild can read it; the one
    part they most need is the narrator speaking.

    Pairing is BY SECTION ID and BY ITEM INDEX, matching how the prose
    above is paired. A translated section that is missing, or shorter
    than its source, simply contributes no Spanish line for that item --
    it never shifts the pairing, which would attach one story's
    translation to another's text.
    """
    tgt_by_id = {s.id: s for s in (translated.sections or [])}
    for sec in _server_evidence_sections_of(req):
        h = doc.add_heading(sec.label, level=1)
        try:
            h.runs[0].font.color.rgb = _DARK_BROWN
        except Exception:
            pass
        tgt = tgt_by_id.get(sec.id)
        tgt_items = list(getattr(tgt, "items", None) or [])
        for idx, item in enumerate(sec.items or []):
            doc.add_paragraph(item)
            rendered = tgt_items[idx] if idx < len(tgt_items) else ""
            # Only when the translation actually differs. When the
            # service is unavailable it passes the source text through,
            # and printing the same sentence twice is worse than
            # printing it once.
            if rendered and rendered != item:
                tp = doc.add_paragraph(rendered)
                if tp.runs:
                    tp.runs[0].font.italic = True
                    tp.runs[0].font.color.rgb = _WARM_GREY
        doc.add_paragraph()  # spacer


def _build_threads_docx(
    req: MemoirExportRequest,
    *,
    render_lang: str = "en",
    resolved_photos: Optional[Dict[str, Path]] = None,
) -> bytes:
    """Build DOCX for threads state: grouped sections with bullet items.

    `render_lang` controls only the chrome (title / subtitle / photos
    heading / arc-roles label). Section content has already been
    translated upstream by _translate_request_content when needed.
    `resolved_photos` maps media_id → server-resolved contained path
    (WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01).
    """
    resolved_photos = resolved_photos or {}
    doc = Document()
    _stamp_source_provenance(doc, req)

    # Title
    title_text = _chrome(render_lang, "threads_title").format(narrator=req.narrator_name)
    title = doc.add_heading(title_text, level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN

    # Subtitle
    sub = doc.add_paragraph(_chrome(render_lang, "threads_subtitle"))
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = _WARM_GREY

    # Arc coverage line (if available) — arc_roles themselves are
    # operator-side labels (display-only), not translated.
    if req.arc_roles:
        arc_line = doc.add_paragraph()
        arc_label = _chrome(render_lang, "story_arc_label").format(
            roles=' · '.join(req.arc_roles),
        )
        arc_run = arc_line.add_run(arc_label)
        arc_run.font.size = Pt(10)
        arc_run.font.color.rgb = _GOLD

    doc.add_paragraph()  # spacer

    # Sections
    for sec in req.sections:
        if not sec.items:
            continue  # skip empty — export truth rule
        h = doc.add_heading(sec.label, level=2)
        h.runs[0].font.color.rgb = _DARK_BROWN

        # Inline photos for this section (Media Builder)
        for photo in _photos_for_section(req, sec.id):
            _add_photo_to_doc(doc, photo, resolved_photos.get(photo.media_id))

        for item in sec.items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

        doc.add_paragraph()  # spacer between sections

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_draft_docx(
    req: MemoirExportRequest,
    *,
    render_lang: str = "en",
    resolved_photos: Optional[Dict[str, Path]] = None,
) -> bytes:
    """Build DOCX for draft state: prose paragraphs, optionally with arc headings.

    `render_lang` controls chrome only; prose content is already
    translated upstream when needed (see _translate_request_content).
    `resolved_photos` maps media_id → server-resolved contained path
    (WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01)."""
    resolved_photos = resolved_photos or {}
    doc = Document()
    _stamp_source_provenance(doc, req)

    title_text = _chrome(render_lang, "draft_title").format(narrator=req.narrator_name)
    title = doc.add_heading(title_text, level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN

    sub = doc.add_paragraph(_chrome(render_lang, "draft_subtitle"))
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = _WARM_GREY

    doc.add_paragraph()  # spacer

    # Build a map of section_key → photos for quick lookup in arc-label detection
    # We use section keys stored on the photo; for draft, we try to match arc labels
    # to memoir section ids (best-effort — draft state doesn't have structured sections).
    section_photos_by_key: dict = {}
    for photo in req.attached_photos:
        section_photos_by_key.setdefault(photo.section_key, []).append(photo)

    if req.prose:
        paragraphs = [p.strip() for p in req.prose.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            lines = para_text.split("\n")
            # Detect arc label marker: "-- Label --"
            if lines and lines[0].strip().startswith("--") and lines[0].strip().endswith("--"):
                label = lines[0].strip().strip("-").strip()
                h = doc.add_heading(label, level=2)
                h.runs[0].font.color.rgb = _DARK_BROWN
                body = "\n".join(lines[1:]).strip()
                if body:
                    doc.add_paragraph(body)
            else:
                doc.add_paragraph(para_text)
            doc.add_paragraph()  # spacer

    # Reviewed captured stories, after the operator's own prose. 2026-08-19.
    _render_evidence_into_draft(doc, req)

    # Append photo section at end of draft (no per-section matching in pure prose)
    # Only include photos not already displayed via section matching
    if req.attached_photos:
        doc.add_page_break()
        ph = doc.add_heading(_chrome(render_lang, "photos_heading"), level=1)
        ph.runs[0].font.color.rgb = _DARK_BROWN
        for photo in req.attached_photos:
            _add_photo_to_doc(doc, photo, resolved_photos.get(photo.media_id))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── WO-ML-04 Phase 4B — bilingual builder ────────────────────────────────────

def _build_threads_docx_bilingual(
    req: MemoirExportRequest,
    translated: MemoirExportRequest,
    *,
    resolved_photos: Optional[Dict[str, Path]] = None,
) -> bytes:
    """Build DOCX with English + Spanish content interleaved per
    section. The narrator's section in source language renders first,
    immediately followed by the translation as a quoted block. A
    bilingual reader can follow either side; a Spanish-only reader
    skips the English block.

    `req` is the original (source) request; `translated` is the
    translated copy produced by _translate_request_content.
    """
    resolved_photos = resolved_photos or {}
    doc = Document()
    _stamp_source_provenance(doc, req)

    # Bilingual title pairs — render both languages stacked.
    src_lang = (req.source_language or "en").strip().lower() or "en"
    tgt_lang = (translated.target_language or "es").strip().lower() or "es"

    src_title = _chrome(src_lang, "threads_title").format(narrator=req.narrator_name)
    tgt_title = _chrome(tgt_lang, "threads_title").format(narrator=req.narrator_name)
    title = doc.add_heading(src_title, level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN
    sub_title = doc.add_heading(tgt_title, level=1)
    sub_title.runs[0].font.color.rgb = _WARM_GREY

    src_sub = doc.add_paragraph(_chrome(src_lang, "threads_subtitle"))
    src_sub.runs[0].font.italic = True
    src_sub.runs[0].font.color.rgb = _WARM_GREY
    tgt_sub = doc.add_paragraph(_chrome(tgt_lang, "threads_subtitle"))
    tgt_sub.runs[0].font.italic = True
    tgt_sub.runs[0].font.color.rgb = _GOLD

    if req.arc_roles:
        arc_line = doc.add_paragraph()
        arc_label = _chrome(src_lang, "story_arc_label").format(
            roles=' · '.join(req.arc_roles),
        )
        arc_run = arc_line.add_run(arc_label)
        arc_run.font.size = Pt(10)
        arc_run.font.color.rgb = _GOLD

    doc.add_paragraph()  # spacer

    # Section-by-section: render source items, then translated items
    # in italic. Photos render between source and translation so the
    # spatial relationship reads naturally regardless of which language
    # the reader follows.
    src_sections_by_id = {s.id: s for s in req.sections}
    for tsec in translated.sections:
        ssec = src_sections_by_id.get(tsec.id)
        if ssec is None or not ssec.items:
            continue

        # Header pair: source label as H2, translated label as H3
        h = doc.add_heading(ssec.label, level=2)
        h.runs[0].font.color.rgb = _DARK_BROWN
        if tsec.label and tsec.label.strip() and tsec.label != ssec.label:
            h2 = doc.add_heading(tsec.label, level=3)
            h2.runs[0].font.color.rgb = _WARM_GREY
            h2.runs[0].font.italic = True

        # Inline photos for this section
        for photo in _photos_for_section(req, ssec.id):
            _add_photo_to_doc(doc, photo, resolved_photos.get(photo.media_id))

        # Source items (English)
        for item in ssec.items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

        # Translated items immediately below — italic so the eye can
        # tell which is which without flipping pages.
        for item in tsec.items or []:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item)
            r.font.italic = True
            r.font.color.rgb = _WARM_GREY

        doc.add_paragraph()  # spacer between sections

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_draft_docx_bilingual(
    req: MemoirExportRequest,
    translated: MemoirExportRequest,
    *,
    resolved_photos: Optional[Dict[str, Path]] = None,
) -> bytes:
    """Bilingual variant of _build_draft_docx. Source paragraph,
    then translated paragraph in italic, repeated for each prose
    paragraph the narrator wrote."""
    resolved_photos = resolved_photos or {}
    doc = Document()
    _stamp_source_provenance(doc, req)

    src_lang = (req.source_language or "en").strip().lower() or "en"
    tgt_lang = (translated.target_language or "es").strip().lower() or "es"

    src_title = _chrome(src_lang, "draft_title").format(narrator=req.narrator_name)
    tgt_title = _chrome(tgt_lang, "draft_title").format(narrator=req.narrator_name)
    title = doc.add_heading(src_title, level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN
    sub_title = doc.add_heading(tgt_title, level=1)
    sub_title.runs[0].font.color.rgb = _WARM_GREY

    src_sub = doc.add_paragraph(_chrome(src_lang, "draft_subtitle"))
    src_sub.runs[0].font.italic = True
    src_sub.runs[0].font.color.rgb = _WARM_GREY
    tgt_sub = doc.add_paragraph(_chrome(tgt_lang, "draft_subtitle"))
    tgt_sub.runs[0].font.italic = True
    tgt_sub.runs[0].font.color.rgb = _GOLD

    doc.add_paragraph()  # spacer

    src_paragraphs = (req.prose or "").split("\n\n") if req.prose else []
    tgt_paragraphs = (translated.prose or "").split("\n\n") if translated.prose else []

    # Pair source + translated paragraphs index-aligned. If the
    # translation pass produced a different paragraph count (rare —
    # paragraph-by-paragraph translation preserves count), pad the
    # shorter list with empty strings.
    n = max(len(src_paragraphs), len(tgt_paragraphs))
    for i in range(n):
        s = src_paragraphs[i] if i < len(src_paragraphs) else ""
        t = tgt_paragraphs[i] if i < len(tgt_paragraphs) else ""
        if not s and not t:
            continue
        s_lines = s.split("\n") if s else []
        # Detect arc-label marker on the source side, mirror to translation
        if s_lines and s_lines[0].strip().startswith("--") and s_lines[0].strip().endswith("--"):
            label = s_lines[0].strip().strip("-").strip()
            h = doc.add_heading(label, level=2)
            h.runs[0].font.color.rgb = _DARK_BROWN
            body = "\n".join(s_lines[1:]).strip()
            if body:
                doc.add_paragraph(body)
            if t:
                tp = doc.add_paragraph(t)
                if tp.runs:
                    tp.runs[0].font.italic = True
                    tp.runs[0].font.color.rgb = _WARM_GREY
        else:
            if s:
                doc.add_paragraph(s)
            if t:
                tp = doc.add_paragraph(t)
                if tp.runs:
                    tp.runs[0].font.italic = True
                    tp.runs[0].font.color.rgb = _WARM_GREY
        doc.add_paragraph()  # spacer

    # Server-owned evidence reaches the BILINGUAL draft in BOTH
    # languages -- see `_render_evidence_into_bilingual_draft` for why
    # source-only was the wrong answer. 2026-08-19.
    _render_evidence_into_bilingual_draft(doc, req, translated)

    if req.attached_photos:
        doc.add_page_break()
        ph = doc.add_heading(
            _chrome(src_lang, "photos_heading") + " · " + _chrome(tgt_lang, "photos_heading"),
            level=1,
        )
        ph.runs[0].font.color.rgb = _DARK_BROWN
        for photo in req.attached_photos:
            _add_photo_to_doc(doc, photo, resolved_photos.get(photo.media_id))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Route ─────────────────────────────────────────────────────────────────────

# ── WO-MEMOIR-STORY-CANDIDATES-WIRE-01: captured-story harvest ────────────

_MEMOIR_ERA_ORDER = [
    "earliest_years", "early_school_years", "adolescence",
    "coming_of_age", "building_years", "later_years", "today",
]


#: Section ids this server OWNS. A client may send operator-authored
#: sections freely -- that is the editing surface doing its job -- but it
#: may not send one wearing a reserved id, because the memoir would then
#: contain a chapter of "captured stories" that no review ever cleared,
#: indistinguishable in the finished document from ones that were.
_RESERVED_STORY_SECTION_PREFIX = "captured_stories"

#: Every section id namespace the SERVER owns. Both lanes are harvested
#: server-side and review-gated -- captured stories by `review_status`,
#: trip stories by the operator's explicit `include_in_memoir` -- so a
#: client section wearing either prefix would appear in the finished
#: document as cleared evidence beside the real thing, and could collide
#: with a server section id (the bilingual builder keys by id, last wins).
#: `trip_stories` was missing from this defence until 2026-08-19.
_RESERVED_SECTION_PREFIXES = ("captured_stories", "trip_stories")


def _is_server_evidence_section(section: Any) -> bool:
    sid = str(getattr(section, "id", "") or "")
    return sid.startswith(_RESERVED_SECTION_PREFIXES)


def _story_source_digest(candidate_id: str) -> str:
    """A stable, non-identifying marker for one captured story.

    A raw narrator UUID must not appear in a document a family reads, and
    a memoir with no provenance at all cannot be traced back to the turn
    it came from -- the audit found exported prose was completely flat.
    A short digest of the candidate id is stable across exports, reveals
    nothing on its face, and lets an operator match a paragraph to its
    source by digesting the id again.
    """
    from ..services.memoir_contract import story_source_id as _sid
    return _sid(candidate_id)


def _captured_story_sections(person_id: str) -> Tuple[List[MemoirSection], str]:
    """Harvest reviewed story candidates into era-grouped memoir sections.

    Rewritten 2026-08-19 (WO-LORI-CONVERSATION-TO-LIFE-MAP-MEMOIR-01
    Commit 3). Returns `(sections, status)` where status is "read" or the
    projection's own failure verdict, because the caller must be able to
    tell an empty narrator from an unreadable one -- see the export route.

    VERBATIM narrator transcripts. The whole point of the preservation
    lane is the narrator's own words: no summarising, no rewriting.

    ── WHAT CHANGED, AND WHY IT MATTERED ───────────────────────────────

    This function used to read `story_candidate_list_for_memoir` and then
    interpret `era_candidates[0]` itself:

        era = eras[0] if eras else "_unplaced"

    That is a second, independent reading of placement, and it disagreed
    with the canonical one. `story_projection` had already decided that
    an era candidate nobody confirmed is NOT a placement; the memoir filed
    the story under it anyway. A machine guess became a chapter heading
    in a document a family keeps, with nothing on the page to say it was
    a guess.

    Placement now comes from `story_projection.memoir_projection`, which
    is the same service the Life Map, the chronology and Lori read. An
    unplaced story -- including one with a year but no era -- goes to
    "More stories", and no era is ever derived for it.

    EACH ELIGIBLE CANDIDATE APPEARS EXACTLY ONCE, keyed by candidate id.
    Deliberately NOT deduplicated by text: two tellings of the same
    memory are two things the narrator said, and collapsing them would be
    the system deciding which of a person's own words to discard.
    """
    try:
        from ..services import story_projection as _sp
        projection = _sp.memoir_projection(person_id)
    except Exception as exc:
        logger.warning("[memoir-docx] story harvest failed: %s", exc)
        return [], "unavailable"
    if not projection.available:
        return [], projection.status
    if not projection.items:
        return [], "read"

    try:
        from ..lv_eras import era_id_to_warm_label as _warm
    except Exception:
        _warm = lambda e: e  # noqa: E731

    seen: set = set()
    by_era: Dict[str, List[Dict[str, Any]]] = {}
    for row in projection.items:
        cid = row.get("id")
        if not cid or cid in seen:
            continue
        seen.add(cid)
        era = row.get("era") or "_unplaced"
        by_era.setdefault(era, []).append(row)

    def _section(era_key: str, label: str, rows: List[Dict[str, Any]]):
        return MemoirSection(
            id=f"{_RESERVED_STORY_SECTION_PREFIX}_{era_key}",
            label=label,
            items=[r["transcript"] for r in rows],
            sources=[_story_source_digest(r["id"]) for r in rows],
            languages=[r.get("language") or "en" for r in rows],
        )

    sections: List[MemoirSection] = []
    ordered = [e for e in _MEMOIR_ERA_ORDER if e in by_era]
    ordered += [e for e in by_era
                if e not in _MEMOIR_ERA_ORDER and e != "_unplaced"]
    for era in ordered:
        try:
            label = "In their own words — " + str(_warm(era))
        except Exception:
            label = "In their own words"
        sections.append(_section(era, label, by_era[era]))
    if "_unplaced" in by_era:
        sections.append(_section(
            "more", "In their own words — More stories", by_era["_unplaced"]))
    return sections, "read"


def _trip_note_source_digest(note_id: str) -> str:
    """Provenance for one approved trip note, from its durable id.

    Trip notes DO have a stable id (`trip_location_notes.id`), so their
    provenance is real rather than invented -- the earlier position that
    they must carry none was over-cautious. The digest is namespaced
    apart from story candidates so the two lanes can never collide.
    """
    from ..services.memoir_contract import trip_note_source_id as _nid
    return _nid(note_id)


def _trip_story_sections(person_id: str) -> Tuple[List[MemoirSection], str]:
    """Harvest APPROVED Travel Doc trip stories into memoir sections.

    WO-MEMOIR-TRIP-STORY-LANE-01 (2026-07-27). Travel Doc modal turns are
    captured to trip_location_notes (source_surface=travel_doc_modal) and
    are deliberately NEVER written to the narrator's life-story archive --
    the two-surface rule of 2026-07-09, locked by
    tests/test_modal_archive_boundary.py. This lane is the sanctioned way
    trip material reaches the memoir: a DB read, gated on the operator's
    explicit include_in_memoir=1, rendered as its own clearly-sourced
    section. It performs no archive write of any kind.

    ── IT NO LONGER FAILS OPEN, 2026-08-19 ─────────────────────────────

    This returned a bare list and swallowed every failure: an unreadable
    trip list produced `[]`, and a single unreadable trip was `continue`d
    past. Both produced a memoir that LOOKED complete while missing
    approved material, which a family cannot detect -- they simply never
    see the chapter.

    It now returns `(sections, status)`:

        read           every requested trip was read, and there is
                       something to show
        empty          read successfully; this narrator has no approved
                       trip notes. A separate word from `read` because a
                       caller that logs the status should be able to say
                       which happened without also counting sections.
        not_attempted  the trips feature is off for this deployment
        unavailable    the trip list itself could not be read
        partial        some trips read, at least one did not

    `partial` is the one worth having separately: it is the case that
    used to be silently `continue`d, and it is the case where the
    document is most convincingly wrong.
    """
    if os.getenv("HORNELORE_TRIPS", "0").strip().lower() not in (
        "1", "true", "yes", "on",
    ):
        return [], "not_attempted"
    try:
        from ..services import trip_repository as _tr
        trips = _tr.trip_list(person_id)
    except Exception as exc:
        logger.warning("[memoir-docx] trip harvest failed: %s", exc)
        return [], "unavailable"
    if not trips:
        return [], "empty"

    def _order(t: Dict[str, Any]):
        """Dated trips in chronological order, undated ones after."""
        start = (t.get("start_date") or "").strip()
        return (0, start) if start else (1, (t.get("created_at") or ""))

    sections: List[MemoirSection] = []
    unreadable = 0
    for trip in sorted(trips, key=_order):
        trip_id = trip.get("id")
        if not trip_id:
            continue
        try:
            notes = _tr.location_notes_list(trip_id)
        except Exception as exc:
            logger.warning(
                "[memoir-docx] trip notes unreadable trip=%s: %s",
                trip_id, exc)
            unreadable += 1
            continue
        items: List[str] = []
        digests: List[str] = []
        languages: List[str] = []
        seen_notes: set = set()
        for n in notes:
            if not n.get("include_in_memoir"):
                continue          # unapproved never reaches the memoir
            note_id = n.get("id")
            if not note_id or note_id in seen_notes:
                continue          # exactly once, by durable note id
            text = (n.get("note_text") or "").strip()
            if not text:
                continue
            seen_notes.add(note_id)
            title = (n.get("note_title") or "").strip()
            items.append(f"{title} \u2014 {text}" if title else text)
            digests.append(_trip_note_source_digest(note_id))
            # THE COLUMN IS `target_language`, and the name is worth a
            # sentence. Verified against the live schema on 2026-08-19:
            # `trip_location_notes` has no `language` column at all -- the
            # first cut read one, so every note silently defaulted to
            # English. `target_language` is set once at note creation,
            # defaults to 'en', and nothing translates a note afterwards,
            # so in practice it records the language the note text IS in.
            # That is how it is read here. No new field is invented; if
            # the intent ever diverges from the usage, this is the line
            # that has to change.
            languages.append(
                str(n.get("target_language") or "").strip().lower() or "en")
        if not items:
            continue
        # PRECEDENCE. This was `"From your travels — " + str(...).strip()
        # or "From your travels"`, and `+` binds tighter than `or`, so an
        # untitled trip produced the dangling `"From your travels — "`
        # rather than the intended fallback.
        _title = str(trip.get("title") or "").strip()
        label = ("From your travels \u2014 " + _title) if _title \
            else "From your travels"
        sections.append(MemoirSection(
            id=f"trip_stories_{trip_id}", label=label, items=items,
            sources=digests, languages=languages))

    if unreadable:
        return sections, "partial"
    return sections, ("read" if sections else "empty")

def _safe_filename_component(raw: Optional[str], *, fallback: str, max_len: int = 80) -> str:
    cleaned = "".join(
        c if (c.isascii() and (c.isalnum() or c in "-_.")) else "_"
        for c in (raw or "")
    )[:max_len].strip("_.")
    return cleaned or fallback


@router.post("/export-docx")
def api_memoir_export_docx(req: MemoirExportRequest):
    """
    Accept memoir content JSON, return a DOCX file as a streaming download.
    Called by memoirExportDOCX() in hornelore1.0.html.

    WO-ML-04 / Phase 4B (2026-05-07): the route now dispatches by
    target_language. 'en' (default) is byte-stable. 'es' translates
    every section item + prose paragraph via services.translation
    before rendering. 'bilingual' renders English + Spanish
    interleaved per section / paragraph.

    WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5):
      - gated behind HORNELORE_MEMOIR_EXPORT_ENABLED (404 when off);
      - person_id must exist in people (422 otherwise);
      - attached photos resolve server-side via the media table and
        must stay contained in the media root (422 on any failure);
      - Content-Disposition filename is allowlist-sanitized.
    """
    _require_enabled()

    if not _DOCX_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="python-docx is not installed on this server. Install with: pip install python-docx",
        )

    # WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5.4): a
    # supplied person_id must name a real narrator before it can scope
    # captured-story harvest or media ownership checks.
    if req.person_id:
        from .. import db as _db
        if not _db.get_person(req.person_id):
            raise HTTPException(
                status_code=422,
                detail=f"person_id '{req.person_id}' not found in people",
            )

    # WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5.3):
    # resolve every attached photo through the media authority BEFORE
    # rendering. Containment/authority failures are loud 422s; only the
    # server-resolved paths below ever reach doc.add_picture().
    resolved_photos = _resolve_attached_photos(req)

    target_lang = _normalize_target_lang(req)

    # ── CLIENT SECTIONS ARE SANITISED UNCONDITIONALLY ──────────────────
    #
    # 2026-08-19. This block used to sit inside
    # `if req.person_id and req.include_captured_stories`, so a caller
    # could bypass the whole defence by omitting the narrator or setting
    # `include_captured_stories=false` -- and then send sections wearing
    # `captured_stories_*` or `trip_stories_*` ids, carrying forged
    # `sources` digests, which the artifact would present as reviewed
    # evidence with server provenance. A sanitiser you can switch off by
    # asking for less is not a sanitiser.
    #
    # Both reserved namespaces are stripped, and every client-supplied
    # `sources`/`languages` array is discarded, on EVERY request. Client
    # CONTENT is untouched: operator-authored prose is the editing
    # surface doing its job and is passed through exactly as sent.
    _client_sections = [
        s.model_copy(update={"sources": [], "languages": []})
        if hasattr(s, "model_copy") else s
        for s in (req.sections or [])
        if not _is_server_evidence_section(s)
    ]
    _spoofed = len(req.sections or []) - len(_client_sections)
    if _spoofed:
        logger.warning(
            "[memoir-docx] dropped %d client section(s) using a reserved "
            "%r id namespace — server evidence is harvested and "
            "review-gated, never accepted from the wire",
            _spoofed, _RESERVED_SECTION_PREFIXES)
    req = req.model_copy(update={"sections": _client_sections}) \
        if hasattr(req, "model_copy") else req
    if not hasattr(req, "model_copy"):
        req.sections = _client_sections

    # ── AUTHORITATIVE LANES: READ, OR REFUSE ───────────────────────────
    #
    # Each requested lane reports read / empty / not_attempted / partial /
    # unavailable. `partial` and `unavailable` REFUSE the export.
    #
    # The alternative -- exporting what was readable -- produces a
    # document that looks complete and is not, and the one person who
    # could notice is the narrator, who will never see the chapter that
    # is missing. A refusal is recoverable; a plausible gap is not.
    _server_sections: List[MemoirSection] = []
    _lane_status: Dict[str, str] = {}

    if req.person_id and req.include_captured_stories:
        _story_sections, _story_status = _captured_story_sections(req.person_id)
        _lane_status["captured_stories"] = _story_status
        if _story_status != "read":
            raise HTTPException(
                status_code=503,
                detail=("reviewed stories could not be read (%s) — export "
                        "refused rather than produce a memoir that looks "
                        "complete" % _story_status))
        _server_sections += _story_sections
        if _story_sections:
            logger.info(
                "[memoir-docx] captured stories appended: %d section(s), "
                "%d stor%s",
                len(_story_sections),
                sum(len(s.items) for s in _story_sections),
                "y" if sum(len(s.items) for s in _story_sections) == 1 else "ies",
            )

    if req.person_id and req.include_trip_stories:
        _trip_sections, _trip_status = _trip_story_sections(req.person_id)
        _lane_status["trip_stories"] = _trip_status
        # `not_attempted` is the trips feature being off for this
        # deployment, which is a configuration answer rather than a
        # failure, so it does not refuse.
        if _trip_status in ("partial", "unavailable"):
            raise HTTPException(
                status_code=503,
                detail=("approved trip stories could not be fully read (%s) "
                        "— export refused rather than produce a memoir that "
                        "looks complete" % _trip_status))
        _server_sections += _trip_sections
        if _trip_sections:
            logger.info(
                "[memoir-docx] trip stories appended: %d section(s), %d note(s)",
                len(_trip_sections),
                sum(len(s.items) for s in _trip_sections))

    if _server_sections:
        # One digest per item, checked rather than assumed. A short or
        # long `sources` array means the positional provenance mapping
        # would attribute a paragraph to the wrong candidate, which is
        # worse than having none.
        for sec in _server_sections:
            # EXACT equality, including an EMPTY array. `sec.languages and`
            # let a server section with no language metadata through, and
            # that is the case where the translator is then handed every
            # item as the request-level language -- the defect this whole
            # field exists to stop. Server evidence always knows its
            # language; a section that has lost it is broken, not lenient.
            if len(sec.languages) != len(sec.items):
                logger.error(
                    "[memoir-docx] language array misaligned section=%s "
                    "items=%d languages=%d", sec.id, len(sec.items),
                    len(sec.languages))
                raise HTTPException(
                    status_code=500,
                    detail="server evidence language metadata is misaligned "
                           "— export refused")
            if len(sec.sources) != len(sec.items):
                logger.error(
                    "[memoir-docx] provenance misaligned section=%s items=%d "
                    "sources=%d", sec.id, len(sec.items), len(sec.sources))
                raise HTTPException(
                    status_code=500,
                    detail="server evidence provenance is misaligned — "
                           "export refused")
        req = req.model_copy(update={
            "sections": list(req.sections) + _server_sections,
        }) if hasattr(req, "model_copy") else req
        if not hasattr(req, "model_copy"):
            req.sections = list(req.sections) + _server_sections

    logger.info(
        "[memoir-docx] export narrator=%s state=%s src=%s tgt=%s sections=%d photos=%d prose_len=%d",
        req.narrator_name, req.memoir_state, req.source_language, target_lang,
        len(req.sections), len(req.attached_photos),
        len(req.prose or ""),
    )

    # WO-POST-REVIEW-SAFETY-DRAFT-EXPORT-HARDENING-01 (Phase 5.5):
    # allowlist-sanitize BOTH client-supplied filename components so no
    # CR/LF/quote/slash/control char can reach the header.
    safe_name = _safe_filename_component(
        (req.narrator_name or "").strip().lower().replace(" ", "_"),
        fallback="memoir",
    )
    safe_state = _safe_filename_component(req.memoir_state, fallback="threads", max_len=20)
    # Filename suffix carries language so re-exports don't overwrite
    # each other when the operator iterates en → es → bilingual.
    lang_suffix = "" if target_lang == "en" else f"_{target_lang}"
    filename = f"lorevox_memoir_{safe_name}_{safe_state}{lang_suffix}.docx"

    # Dispatch by target language.
    if target_lang == "en":
        # Pre-Phase-4B path. Byte-stable.
        if req.memoir_state == "draft":
            docx_bytes = _build_draft_docx(
                req, render_lang="en", resolved_photos=resolved_photos)
        else:
            docx_bytes = _build_threads_docx(
                req, render_lang="en", resolved_photos=resolved_photos)
    elif target_lang == "es":
        # Translate first, then render with Spanish chrome.
        translated = _translate_request_content(req, "es")
        _assert_translation_covered(req, translated, "es")
        if req.memoir_state == "draft":
            docx_bytes = _build_draft_docx(
                translated, render_lang="es", resolved_photos=resolved_photos)
        else:
            docx_bytes = _build_threads_docx(
                translated, render_lang="es", resolved_photos=resolved_photos)
    else:  # bilingual
        # Translate to Spanish; render with both languages interleaved.
        # Source language defaults to 'en' for the v1 scope; future
        # work can extend bilingual to other source languages.
        translated = _translate_request_content(req, "es")
        # Bilingual was NOT checked. The bilingual builder suppresses an
        # identical second paragraph, so a failed translation produced a
        # source-only document that looked deliberate -- the same silent
        # loss the Spanish path already refused.
        _assert_translation_covered(req, translated, "es")
        if req.memoir_state == "draft":
            docx_bytes = _build_draft_docx_bilingual(
                req, translated, resolved_photos=resolved_photos)
        else:
            docx_bytes = _build_threads_docx_bilingual(
                req, translated, resolved_photos=resolved_photos)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/canonical")
def api_memoir_canonical(person_id: str):
    """The reviewed evidence this narrator's memoir will contain.

    WO-LORI-CONVERSATION-TO-LIFE-MAP-MEMOIR-01 Commit B (2026-08-19).

    THE POINT OF THIS ROUTE IS THAT THE OPERATOR CAN SEE IT BEFORE
    EXPORTING. Reviewed stories and approved trip notes used to be
    appended server-side during DOCX generation and nowhere else, so the
    preview and the TXT export showed a document the DOCX did not match.
    The panel now reads this, the TXT export renders this, and the DOCX
    harvests the same lanes -- and every item carries a `source_id` so
    "exactly once" is checkable across all three rather than hoped for.

    Read-only. Approves nothing, places nothing, writes nothing.
    """
    _require_enabled()
    if not (person_id or "").strip():
        raise HTTPException(status_code=422, detail="person_id required")
    from .. import db as _db
    if not _db.get_person(person_id):
        raise HTTPException(
            status_code=422,
            detail=f"person_id '{person_id}' not found in people")
    from ..services.memoir_contract import canonical_memoir
    return canonical_memoir(person_id).as_dict()
