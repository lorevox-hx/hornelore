"""Trip memoir DOCX builder — WO-TRIP-IMPORT-AND-CLUSTER-01 Phase 4.

Deterministic dual-axis render per WO-TRIP-MEMOIR-01 (no LLM
authoring): Part I walks regions -> nested stops in order, Part II
walks themes with their matching stops, Part III embeds the clustered
photo appendix (include_in_memoir=1 links joined to photos.image_path).

python-docx is imported lazily and guarded — callers get a clear
RuntimeError when it isn't installed (the memoir_export router uses
the same posture with a 503).

WO-EVIDENCE-LIFECYCLE-TRIP-FORCE-01: hidden evidence (notes / sources
/ photo links stamped hidden=1) is excluded UPSTREAM — the preview
dict (trip_repository.trip_memoir_preview) and the photo rows
(photo_links_with_photo_paths) this builder renders are assembled from
hide-aware repository reads, so a hidden row never reaches this
module regardless of its include_in_memoir flag.
"""
from __future__ import annotations

import io
import logging
import os
from typing import Any, Dict, List, Optional

logger = logging.getLogger("code.api.services.trip_memoir_docx")


# WO-TRAVEL-DOC-CLOSEOUT-01 — `_safe_caption` MOVED, not deleted.
#
# It lived here and chose the caption for the document. The browser
# preview had no equivalent, so the operator reviewed a count while the
# family received captions. The rule now lives beside the grouping, in
# `trip_repository._safe_photo_caption`, and BOTH consumers read the one
# projection that applies it.
#
# Deliberately not left here as a second copy. Two implementations of
# "which caption is safe to print" would drift, and the direction they
# drift in is unapproved text reaching a family document -- which has
# already happened once on this path.


def build_trip_docx(
    preview: Dict[str, Any],
    photo_rows: Optional[List[Dict[str, Any]]] = None,
    appendix: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Render the memoir-preview dict + the photo-appendix projection.

    `appendix` is the projection the caller already built, so an export
    reads the photo-link table ONCE and the document describes exactly
    the appendix the operator reviewed. `photo_rows` remains for callers
    that have rows and no projection; it is built from them.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception as exc:  # pragma: no cover - environment-dependent
        raise RuntimeError(
            "python-docx is not installed on this server. "
            "Install with: pip install python-docx"
        ) from exc

    doc = Document()

    def _story_notes(notes, indent=0.0):
        # Promoted story notes (include_in_memoir=1). Empty stays empty —
        # no invented prose.
        for note in notes or []:
            t = (note.get("note_title") or "").strip()
            body = (note.get("note_text") or "").strip()
            if t:
                h = doc.add_paragraph()
                r = h.add_run(t); r.bold = True; r.font.size = Pt(10)
                if indent:
                    h.paragraph_format.left_indent = Inches(indent)
            if body:
                pn = doc.add_paragraph(body)
                pn.runs[0].font.size = Pt(10)
                if indent:
                    pn.paragraph_format.left_indent = Inches(indent)

    def _sources(srcs, indent=0.0):
        # Promoted sources (include_in_memoir=1) as a compact Sources block.
        for s in srcs or []:
            label = (s.get("title") or s.get("filename") or
                     s.get("source_type") or "Source").strip()
            h = doc.add_paragraph()
            r = h.add_run("Source — " + label); r.bold = True; r.font.size = Pt(9)
            if indent:
                h.paragraph_format.left_indent = Inches(indent)
            detail = (s.get("summary") or s.get("pasted_text") or
                      s.get("link_url") or "").strip()
            if detail:
                pn = doc.add_paragraph(detail)
                pn.runs[0].font.size = Pt(9)
                if indent:
                    pn.paragraph_format.left_indent = Inches(indent)

    title = preview.get("title") or "Trip Memoir"
    dr = preview.get("date_range") or {}
    doc.add_heading(str(title), level=0)
    date_line = " — ".join(
        [d for d in (dr.get("start"), dr.get("end")) if d]
    )
    if date_line:
        p = doc.add_paragraph(date_line)
        p.runs[0].font.size = Pt(11)
    if preview.get("summary"):
        doc.add_paragraph(str(preview["summary"]))
    _story_notes(preview.get("story_notes"))
    _sources(preview.get("sources"))

    # [`appendix` / `photo_rows` built the Part III photo appendix and
    # the per-stop "· N approved photos" line. Both are retired: the
    # timeline prints each photograph once under its own day, so there
    # is no appendix to describe and no approval to count. The
    # parameters are still accepted, and ignored, so an existing caller
    # that passes them keeps working rather than raising.]

    # ── Part I — The Journey in Order ────────────────────────────────
    doc.add_heading("Part I — The Journey in Order", level=1)

    def _stop_paragraph(stop: Dict[str, Any], depth: int) -> None:
        name = stop.get("title") or stop.get("location_name") or ""
        bits = [str(name)]
        ds, de = stop.get("date_start"), stop.get("date_end")
        if ds and de and ds != de:
            bits.append(f"({ds} – {de})")
        elif ds:
            bits.append(f"({ds})")
        stype = stop.get("stop_type")
        if stype and stype not in ("sight",):
            bits.append(f"[{stype}]")
        # [Read "· N approved photos" from the appendix projection.
        # Retired 2026-08-06 with the appendix: a count of approved
        # photographs, in a document where approval decides nothing and
        # the photographs themselves print under their day, pointed at
        # a section that no longer exists.]
        style = "List Bullet" if depth == 0 else "List Bullet 2"
        doc.add_paragraph(" ".join(bits), style=style)
        if stop.get("notes"):
            note = doc.add_paragraph(str(stop["notes"]))
            note.paragraph_format.left_indent = Inches(0.5 + 0.25 * depth)
            note.runs[0].font.size = Pt(10)
        _story_notes(stop.get("story_notes"), indent=0.5 + 0.25 * depth)
        _sources(stop.get("sources"), indent=0.5 + 0.25 * depth)
        for child in stop.get("day_trips", []):
            _stop_paragraph(child, depth + 1)

    for i, region in enumerate(preview.get("part_one_journey_in_order", []), 1):
        heading_bits = [f"{i}. {region.get('region') or ''}"]
        rdr = region.get("date_range") or {}
        if rdr.get("start"):
            heading_bits.append(f"({rdr.get('start')} – {rdr.get('end') or ''})")
        doc.add_heading(" ".join(heading_bits), level=2)
        if region.get("base_address"):
            doc.add_paragraph(f"Base: {region['base_address']}")
        if region.get("summary"):
            doc.add_paragraph(str(region["summary"]))
        _story_notes(region.get("story_notes"))
        _sources(region.get("sources"))
        for stop in region.get("stops", []):
            _stop_paragraph(stop, 0)

    # ── The timeline — WO-TRAVEL-DOC-CLOSEOUT-01, 2026-08-06 ─────────
    #
    # [This block projected `part_one_days`, which carried only the
    # APPROVED days and only the day card's own six text fields.
    # Retired with that design: the product rule is that the visible
    # trip timeline is the editable source of truth and this document
    # is a snapshot of it, so nothing here filters on approval and
    # nothing here re-interprets a day. `part_one_timeline` is the same
    # projection the operator is looking at on screen.]
    #
    # It renders AFTER the region walk: a trip that has both keeps its
    # planned route first and its lived days second, and a trip with no
    # regions -- the case that produced the empty Part I -- starts here.
    _timeline = preview.get("part_one_timeline") or {}
    _speaker = str(preview.get("narrator_label") or "Narrator").strip() \
        or "Narrator"

    def _photo_caption(item):
        """Caption plus, where it is machine text, a label saying so.

        Rule 8. `caption_source` distinguishes the narrator's own words
        from an operator caption from a description this system
        generated. Flattening them is how a machine sentence comes to
        sit under a family photograph looking like something Chris
        wrote about his own trip.
        """
        cap = str(item.get("caption") or "").strip()
        src = str(item.get("caption_source") or "").strip()
        if cap and src == "machine":
            return "Draft description (machine-written, not reviewed): " + cap
        return cap

    def _render_items(items, indent=0.25):
        for item in items or []:
            kind = str(item.get("kind") or "")
            if kind == "day_text":
                para = doc.add_paragraph()
                r = para.add_run(str(item.get("label") or "") + ": ")
                r.bold = True; r.font.size = Pt(10)
                r2 = para.add_run(str(item.get("text") or ""))
                r2.font.size = Pt(10)
                para.paragraph_format.left_indent = Inches(indent)
            elif kind == "conversation":
                # Rule 5: both speakers, labelled, in one item -- a link
                # row is one exchange, so printing it once with two
                # labels is what makes rule 10 hold here.
                for who, text in ((_speaker, item.get("narrator_said")),
                                  ("Lori", item.get("lori_said"))):
                    body = str(text or "").strip()
                    if not body:
                        continue
                    para = doc.add_paragraph()
                    r = para.add_run(who + ": ")
                    r.bold = True; r.font.size = Pt(10)
                    r2 = para.add_run(body); r2.font.size = Pt(10)
                    para.paragraph_format.left_indent = Inches(indent)
            elif kind == "note":
                _story_notes([{"note_title": item.get("title"),
                               "note_text": item.get("text")}], indent=indent)
            elif kind == "source":
                _sources([{"title": item.get("title"),
                           "source_type": item.get("source_type"),
                           "summary": item.get("summary"),
                           "link_url": item.get("link_url")}], indent=indent)
            elif kind == "photo":
                path = str(item.get("image_path") or "")
                placed = False
                if path and os.path.isfile(path):
                    try:
                        doc.add_picture(path, width=Inches(4.5))
                        placed = True
                    except Exception as exc:
                        logger.warning(
                            "[trip-docx] photo embed failed path=%s: %s",
                            path, exc)
                bits = [b for b in (_photo_caption(item),
                                    str(item.get("at") or "").strip()) if b]
                if not placed:
                    # Said rather than skipped: a photograph the operator
                    # can see on the timeline and cannot find in the
                    # document reads as a lost photograph.
                    bits.append("(photograph could not be found on disk)")
                if bits:
                    cp = doc.add_paragraph(" — ".join(bits))
                    cp.runs[0].font.size = Pt(9)
                    cp.paragraph_format.left_indent = Inches(indent)

    # EVERY projected day, with no second filter here.
    #
    # [Filtered on `items or title`. A day row is visible on the
    # timeline when it exists -- its number and its date are content,
    # and an operator looking at "Day 2 - 15 July" with nothing under it
    # is looking at a real day of the trip. Two filters also meant two
    # definitions of "a day worth printing", in two languages, which is
    # the second-interpretation problem this whole rewrite is about.
    # The projection decides; the builder renders what it is given.]
    _days = list(_timeline.get("days") or [])
    if _days:
        doc.add_heading("Day by day", level=2)
    for day in _days:
        bits = []
        if day.get("day_index") not in (None, ""):
            bits.append(f"Day {day['day_index']}")
        if day.get("date"):
            bits.append(str(day["date"]))
        head = " · ".join(bits)
        if day.get("title"):
            head = f"{head} — {day['title']}" if head else str(day["title"])
        doc.add_heading(head or "A day on this trip", level=3)
        _render_items(day.get("items"))

    # Rule 6. Material with no day is still the operator's material.
    _unplaced = (_timeline.get("unplaced") or {}).get("items") or []
    if _unplaced:
        doc.add_heading("Needs a day", level=2)
        doc.add_paragraph(
            "Recorded on this trip but not yet placed on a day.")
        _render_items(_unplaced)

    # A heading with nothing under it reads as a deletion, so the empty
    # case says what happened instead.
    if (not _days and not _unplaced
            and not preview.get("part_one_journey_in_order")):
        doc.add_paragraph(
            "(Nothing has been recorded on this trip yet.)")

    # ── Part II — Themes That Ran Through the Trip ───────────────────
    doc.add_heading("Part II — Themes That Ran Through the Trip", level=1)
    themes = preview.get("part_two_themes", [])
    if not themes:
        doc.add_paragraph("(No themes recorded for this trip yet.)")
    for theme in themes:
        doc.add_heading(str(theme.get("theme") or ""), level=2)
        if theme.get("description"):
            doc.add_paragraph(str(theme["description"]))
        stops = theme.get("stops") or []
        if stops:
            doc.add_paragraph(
                "Across: " + ", ".join(str(s) for s in stops)
            )

    # ── Part III — Photo Appendix: RETIRED 2026-08-06 ────────────────
    #
    # [This heading embedded every memoir-approved photograph, grouped
    # by stop, region or day, and closed with "(N photos embedded)".]
    #
    # Removed because the timeline now prints each photograph under the
    # day it is placed on, and rule 10 is absolute: every visible
    # timeline item appears exactly once. Keeping the appendix as well
    # would embed the same image twice -- once where it happened and
    # once in a list at the back -- which is both a duplication and, at
    # roughly 5 MB a photograph, a doubling of the file.
    #
    # Rule 11 permits an appendix ("if the appendix remains") on two
    # conditions: it must not be the only place the photograph appears,
    # and it must not contradict the day placement. Printing inline
    # satisfies the first by construction and cannot contradict the
    # second, because the placement IS the heading it sits under.
    #
    # Photographs with no day are not lost: they print under "Needs a
    # day" above, which is where the operator will go to place them.

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
